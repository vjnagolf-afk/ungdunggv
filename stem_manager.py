import streamlit as st
import io
from docx import Document

# Hàm hỗ trợ tạo file Word từ nội dung văn bản
def create_word_file(title, content):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =========================================================
# HÀM HIỂN THỊ THẺ 1: CÁC SẢN PHẨM STEM
# =========================================================
def render_tab_1():
    st.info("💡 **Giao diện Bộ lọc Đề xuất:** Thầy chọn các tiêu chí dưới đây để AI gợi ý danh sách dự án STEM phù hợp.")
    
    col_m1, col_m2, col_m3 = st.columns(3)
    with col_m1:
        chon_khoi_t1 = st.selectbox("1. Khối lớp (Thẻ 1):", ["Khối 9", "Khối 8", "Khối 7", "Khối 6"], key="khoi_t1")
    with col_m2:
        chon_mon_t1 = st.selectbox("2. Môn chủ đạo (Thẻ 1):", ["Khoa học tự nhiên", "Toán học", "Công nghệ", "Tin học"], key="mon_t1")
    with col_m3:
        chon_chu_de_t1 = st.selectbox("3. Lĩnh vực (Thẻ 1):", [
            "Tiết kiệm năng lượng", 
            "Bảo vệ môi trường", 
            "Nông nghiệp thông minh", 
            "Nhà thông minh (Smart Home)", 
            "Chủ đề tự do"
        ], key="chude_t1")
        
    mon_tich_hop_t1 = st.multiselect(
        "Các môn học cần tích hợp (Thẻ 1):", 
        ["Toán học", "Khoa học", "Công nghệ", "Kỹ thuật", "Nghệ thuật"],
        key="montichhop_t1"
    )
    
    tich_hop_ai_t1 = st.checkbox("🔌 Ưu tiên dự án có AI / Vi điều khiển", value=True, key="chk_ai_t1")
    tich_hop_khuyet_tat_t1 = st.checkbox("🤝 Ưu tiên dự án phù hợp Giáo dục hòa nhập", value=True, key="chk_kt_t1")
    
    st.markdown("---")
    
    if st.button("✨ Kích hoạt AI gợi ý danh sách chủ đề (Thẻ 1)", use_container_width=True, key="btn_ai_t1"):
        with st.spinner("AI đang phân tích các tiêu chí và tổng hợp dữ liệu..."):
            st.session_state.stem_ai_suggestions = f"""
            ### 📌 Danh sách dự án STEM đề xuất cho {chon_khoi_t1} - Môn {chon_mon_t1}
            *(Lĩnh vực: {chon_chu_de_t1})*
            
            **1. Hệ thống giám sát và điều khiển {chon_chu_de_t1.lower()} tự động**
            *   **Mô tả:** Sử dụng vi điều khiển ESP8266 kết nối cảm biến để thu thập dữ liệu môi trường.
            *   **Tính hòa nhập:** Cung cấp sơ đồ mạch điện in nổi, phân công vai trò quan sát phù hợp với năng lực học sinh.
            
            **2. Mô hình trực quan ứng dụng AI**
            *   **Mô tả:** Xây dựng mô hình vật lý kết hợp camera để nhận diện, giúp học sinh nắm bắt thực tiễn.
            """
    
    if st.session_state.stem_ai_suggestions:
        with st.container(border=True):
            st.markdown(st.session_state.stem_ai_suggestions)

# =========================================================
# HÀM HIỂN THỊ THẺ 2: XÂY DỰNG DỰ ÁN GIÁO DỤC STEM
# =========================================================
def render_tab_2():
    st.info("🛠️ **Giao diện Thiết kế Kế hoạch Bài dạy:** Thầy nhập thông tin để AI biên soạn KHBD hoàn chỉnh.")
    
    ten_chu_de_t2 = st.text_input("Tên dự án / Chủ đề STEM (Thẻ 2):", 
                               placeholder="Ví dụ: Thiết kế hệ thống tiết kiệm năng lượng trường học...", key="ten_t2")
    
    col1, col2 = st.columns(2)
    with col1:
        mon_chu_dao_t2 = st.selectbox("Môn học chủ đạo (Thẻ 2):", ["Khoa học tự nhiên", "Toán học", "Công nghệ", "Tin học"], key="mon_t2")
        lop_hoc_t2 = st.selectbox("Lớp (Thẻ 2):", ["Lớp 9", "Lớp 8", "Lớp 7", "Lớp 6"], key="lop_t2")
    with col2:
        thoi_luong_t2 = st.text_input("Thời lượng thực hiện (Thẻ 2):", placeholder="Ví dụ: 3 Tiết", key="thoiluong_t2")
    
    tich_hop_ai_iot_t2 = st.checkbox("🔌 Tích hợp ứng dụng AI hoặc Vi điều khiển (VD: ESP8266) trong bài dạy", value=True, key="chk_ai_t2")
    tich_hop_hoa_nhap_t2 = st.checkbox("🤝 Phân hóa câu hỏi và hoạt động cho học sinh khuyết tật", value=True, key="chk_kt_t2")
    
    tai_lieu_nguon = st.file_uploader("Tải lên tài liệu tham khảo (Thẻ 2)", accept_multiple_files=True, key="file_t2")
    
    st.markdown("---")
    
    if st.button("🚀 Kích hoạt AI thiết kế Kế hoạch bài dạy (Thẻ 2)", type="primary", use_container_width=True, key="btn_ai_t2"):
        if not ten_chu_de_t2:
            st.warning("Vui lòng nhập tên chủ đề STEM!")
        else:
            with st.spinner("Hệ thống đang phân tích sư phạm và thiết kế KHBD chi tiết. Quá trình này có thể mất ít phút..."):
                yeu_cau_iot = "Có tích hợp ứng dụng AI hoặc Vi điều khiển (như ESP8266, Arduino) vào sản phẩm STEM." if tich_hop_ai_iot_t2 else "Không bắt buộc tích hợp Vi điều khiển."
                yeu_cau_hoa_nhap = "Bắt buộc có các ghi chú phân hóa, điều chỉnh câu hỏi và hoạt động để hỗ trợ học sinh khuyết tật (giáo dục hòa nhập) trong từng bước tiến trình." if tich_hop_hoa_nhap_t2 else ""

                prompt = f"""
                Đóng vai là một chuyên gia giáo dục STEM và giáo viên cốt cán, hãy soạn một Kế hoạch bài dạy (KHBD) STEM thật chi tiết và chuẩn mực.
                LƯU Ý QUAN TRỌNG: TUYỆT ĐỐI KHÔNG sử dụng từ "Giáo án", CHỈ SỬ DỤNG cụm từ "Kế hoạch bài dạy" hoặc "KHBD".
                
                THÔNG TIN BÀI DẠY:
                - Tên chủ đề/dự án: {ten_chu_de_t2}
                - Môn học chủ đạo: {mon_chu_dao_t2}
                - Đối tượng: {lop_hoc_t2}
                - Thời lượng: {thoi_luong_t2}
                - Yêu cầu kỹ thuật: {yeu_cau_iot}
                - Yêu cầu sư phạm: {yeu_cau_hoa_nhap}
                
                YÊU CẦU CẤU TRÚC CHI TIẾT CỦA KHBD (Soạn đầy đủ các bước, không viết sơ sài, trình bày rõ ràng bằng Markdown):
                
                I. MỤC TIÊU
                1. Năng lực (Năng lực đặc thù môn học, Năng lực chung, Năng lực STEM).
                2. Phẩm chất.
                
                II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU
                1. Giáo viên chuẩn bị.
                2. Học sinh chuẩn bị.
                
                III. TIẾN TRÌNH DẠY HỌC (Trình bày chi tiết Hoạt động của GV, Hoạt động của HS, Sản phẩm dự kiến và Tiêu chí đánh giá cho từng hoạt động theo quy trình 5 bước):
                Hoạt động 1: Xác định vấn đề / Nhu cầu thực tiễn.
                Hoạt động 2: Nghiên cứu kiến thức nền và Đề xuất giải pháp.
                Hoạt động 3: Lựa chọn giải pháp thiết kế (Bản vẽ/Sơ đồ).
                Hoạt động 4: Chế tạo mô hình và Thử nghiệm (Chi tiết cách học sinh sử dụng thiết bị/vi điều khiển nếu có).
                Hoạt động 5: Chia sẻ, Thảo luận và Đánh giá.
                """
                
                try:
                    # Giao diện tạm hiển thị trước khi thầy thay mã API
                    noi_dung_ai = f"""# KẾ HOẠCH BÀI DẠY STEM: {ten_chu_de_t2.upper()}
**Môn học chủ đạo:** {mon_chu_dao_t2} | **Đối tượng:** {lop_hoc_t2} | **Thời lượng:** {thoi_luong_t2}

## I. MỤC TIÊU
**1. Năng lực:**
*   **Năng lực đặc thù:** Học sinh giải thích được nguyên lý hoạt động của mạch điện cơ bản...
*   **Năng lực chung:** Năng lực giải quyết vấn đề và sáng tạo thông qua việc thiết kế...
**2. Phẩm chất:** Chăm chỉ, trách nhiệm trong việc hoàn thiện sản phẩm nhóm...

## II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU
*   **Giáo viên:** Bài trình chiếu, bộ linh kiện vi điều khiển mẫu, bảng tiêu chí đánh giá (Rubric).
*   **Học sinh:** Bìa carton, kéo, keo dán, dây dẫn...

## III. TIẾN TRÌNH DẠY HỌC
### Hoạt động 1: Xác định vấn đề
**a. Mục tiêu:** Kích thích sự tò mò của HS về hiện tượng lãng phí năng lượng...
**b. Tổ chức thực hiện:**
*   **Hoạt động của GV:** Đưa ra tình huống thực tế bằng video...
*   **Hoạt động của HS:** Quan sát, ghi chép và đặt câu hỏi...
*(Giáo dục hòa nhập: Cung cấp phiếu quan sát có hình ảnh cỡ lớn cho HS khiếm thị một phần...)*
"""
                    st.session_state.stem_generated_content = noi_dung_ai
                    st.success("Thiết kế Kế hoạch bài dạy thành công!")
                except Exception as e:
                    st.error(f"Có lỗi xảy ra khi kết nối AI: {e}")

    if st.session_state.stem_generated_content != "":
        st.markdown("### 📖 KẾT QUẢ THIẾT KẾ")
        with st.container(border=True):
            st.markdown(st.session_state.stem_generated_content)
        
        col_download, col_save = st.columns(2)
        with col_download:
            docx_file = create_word_file(ten_chu_de_t2, st.session_state.stem_generated_content)
            st.download_button(
                label="📥 Tải KHBD (File Word)",
                data=docx_file,
                file_name=f"KHBD_STEM_{ten_chu_de_t2.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="btn_down_t2"
            )
        with col_save:
            if st.button("💾 Lưu dự án vào hệ thống (Thẻ 3)", use_container_width=True, key="btn_save_t2"):
                st.session_state.stem_saved_projects[ten_chu_de_t2] = st.session_state.stem_generated_content
                st.toast("Đã lưu KHBD thành công! Hãy kiểm tra ở Thẻ 3.", icon="✅")

# =========================================================
# HÀM HIỂN THỊ THẺ 3: QUẢN LÝ DỰ ÁN ĐÃ LƯU
# =========================================================
def render_tab_3():
    st.info("📁 **Kho Lưu trữ:** Chứa các dự án thầy đã thiết kế và ấn lưu ở Thẻ 2.")
    
    if len(st.session_state.stem_saved_projects) > 0:
        danh_sach_du_an = list(st.session_state.stem_saved_projects.keys())
        for ten_da in danh_sach_du_an:
            with st.expander(f"📌 {ten_da}"):
                st.markdown(st.session_state.stem_saved_projects[ten_da])
                if st.button("🗑️ Xóa dự án này", key=f"btn_del_{ten_da}"):
                    del st.session_state.stem_saved_projects[ten_da]
                    st.rerun() 
    else:
        st.warning("Hiện tại chưa có dự án nào được lưu. Thầy hãy tạo và lưu dự án ở Thẻ 2 nhé.")

# =========================================================
# HÀM CHÍNH ĐƯỢC GỌI TỪ APP.PY
# =========================================================
def render_stem_section():
    st.markdown("## 🚀 HỆ SINH THÁI GIÁO DỤC STEM")
    st.markdown("---")
    
    if "stem_generated_content" not in st.session_state:
        st.session_state.stem_generated_content = ""
    if "stem_saved_projects" not in st.session_state:
        st.session_state.stem_saved_projects = {}
    if "stem_ai_suggestions" not in st.session_state:
        st.session_state.stem_ai_suggestions = ""

    tab1, tab2, tab3 = st.tabs([
        "💡 1. CÁC SẢN PHẨM STEM", 
        "🛠️ 2. XÂY DỰNG DỰ ÁN", 
        "📁 3. QUẢN LÝ DỰ ÁN ĐÃ LƯU"
    ])

    with tab1:
        render_tab_1()
    with tab2:
        render_tab_2()
    with tab3:
        render_tab_3()
