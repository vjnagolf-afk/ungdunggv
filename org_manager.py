import streamlit as st
import sqlite3
import pandas as pd
import io
import re

DB_PATH = "teacher_assistant.db"

def setup_org_database():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS org_members (id INTEGER PRIMARY KEY AUTOINCREMENT, fullname TEXT UNIQUE, position TEXT, main_subject TEXT, email TEXT, phone TEXT, note TEXT)")
    cursor.execute("CREATE TABLE IF NOT EXISTS org_assignments (id INTEGER PRIMARY KEY AUTOINCREMENT, fullname TEXT UNIQUE, subject_class TEXT, homeroom TEXT, concurrent TEXT, total_periods TEXT DEFAULT '0')")
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS org_emulation_years (
        id INTEGER PRIMARY KEY AUTOINCREMENT, nam_hoc TEXT, fullname TEXT, dg_vien_chuc TEXT DEFAULT '', bd_hsg TEXT DEFAULT '',
        nckh TEXT DEFAULT '', stem TEXT DEFAULT '', sang_kien TEXT DEFAULT '', gvdg TEXT DEFAULT '', gvcng TEXT DEFAULT '',
        tdtt_nv TEXT DEFAULT '', hien_mau TEXT DEFAULT '', khac TEXT DEFAULT '', UNIQUE(nam_hoc, fullname)
    )
    """)
    cursor.execute("SELECT COUNT(*) FROM org_members")
    if cursor.fetchone() == 0:
        danh_sach_goc = [
            ("Lê Hồng Dưỡng", "Tổ trưởng", "KHTN (Vật lý) - CN", "vjnagolf@gmail.com", "0984331778", ""),
            ("Nguyễn Thị Huyền Trang", "Tổ viên", "KHTN (Vật lý) - CN", "nthtrangnct@gmail.com", "", ""),
            ("Lý Nguyễn Thu Nhi", "Tổ viên", "KHTN (Vật lý) - CN", "nthtrangnct@gmail.com", "", ""),
            ("Lê Hùng Cường", "Tổ viên", "KHTN (Vật lý) - CN", "nthtrangnct@gmail.com", "", ""),
            ("Khương Thị Thúy Vân", "Tổ viên", "KHTN (Sinh)", "nthtrangnct@gmail.com", "", ""),
            ("Trần Xuân Hạnh", "Tổ viên", "GDTC", "nthtrangnct@gmail.com", "", ""),
            ("Trương Vĩnh Văn", "Tổ viên", "KHTN (Sinh) - GDTC", "nthtrangnct@gmail.com", "", ""),
            ("Phạm Xuân Thọ", "Tổ viên", "KHTN (Sinh) - GDTC", "nthtrangnct@gmail.com", "", ""),
            ("Phạm Thùy Ngoan", "Tổ viên", "KHTN (Hóa)", "", "", ""),
            ("Huỳnh Thị Kim Lý", "Tổ viên", "KHTN", "", "", ""),
            ("Nguyễn Thanh Mai", "Tổ viên", "KHTN", "", "", ""),
            ("Phạm Thị Minh Anh", "Tổ viên", "KHTN", "", "", "")
        ]
        for row in danh_sach_goc:
            cursor.execute("INSERT OR IGNORE INTO org_members (fullname, position, main_subject, email, phone, note) VALUES (?, ?, ?, ?, ?, ?)", row)
            cursor.execute("INSERT OR IGNORE INTO org_assignments (fullname, subject_class, homeroom, concurrent, total_periods) VALUES (?, '-', '-', '-', '0')", (row,))
    conn.commit()
    conn.close()

def clean_dataframe_nan(df):
    df = df.fillna("-")
    for col in df.columns:
        df[col] = df[col].apply(lambda x: "-" if str(x).strip().lower() in ["nan", "none", ""] else str(x).strip())
    return df
def render_org_section():
    setup_org_database()
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🔒 CHỌN VAI TRÒ ĐĂNG NHẬP")
    vai_tro = st.sidebar.selectbox("Vai trò", ["Giáo viên bộ môn", "Tổ trưởng chuyên môn (Admin)", "Ban giám hiệu"], label_visibility="collapsed")
    
    is_admin = False
    if vai_tro == "Tổ trưởng chuyên môn (Admin)":
        ma_pin = st.sidebar.text_input("Nhập mã pin quản lý Admin:", type="password", value="")
        if ma_pin == "123456":
            st.sidebar.success("✅ Quyền Admin đã mở.")
            is_admin = True
        elif ma_pin != "": st.sidebar.error("❌ Mã PIN sai.")

    st.subheader("📋 HỆ THỐNG QUẢN LÝ VÀ PHÂN CÔNG CHUYÊN MÔN GIẢNG DẠY")
    tab1, tab2, tab3 = st.tabs(["👥 Danh sách thành viên", "📊 Phân công giảng dạy", "🏅 Thành tích & Thi đua"])

    with tab1:
        st.markdown("#### 👥 Danh sách thành viên tổ chuyên môn")
        excel_m = st.file_uploader("📥 Tải file Excel danh sách Giáo viên lên hệ thống:", type=["xlsx", "xls"], key="up_m")
        if excel_m and st.button("🚀 Xác nhận nạp danh sách Giáo viên", type="secondary", use_container_width=True):
            try:
                df_up_m = pd.read_excel(excel_m)
                df_up_m.columns = [str(c).strip() for c in df_up_m.columns]
                conn = sqlite3.connect(DB_PATH)
                for _, r in df_up_m.iterrows():
                    conn.execute("INSERT OR REPLACE INTO org_members (fullname, position, main_subject, email, phone, note) VALUES (?, ?, ?, ?, ?, ?)", (str(r.get("Họ và tên", "")), str(r.get("Chức vụ", "Tổ viên")), str(r.get("Phân môn chính", "")), str(r.get("Email", "")), str(r.get("Số điện thoại", "")), str(r.get("Ghi chú", ""))))
                conn.commit(); conn.close()
                st.success("🎉 Đã nạp thành công!"); st.rerun()
            except Exception as e: st.error(f"Lỗi: {e}")

        conn = sqlite3.connect(DB_PATH)
        df_members = pd.read_sql_query("SELECT fullname as [Họ và tên], position as [Chức vụ], main_subject as [Phân môn chính], email as [Email], phone as [Số điện thoại], note as [Ghi chú] FROM org_members", conn)
        conn.close()
        df_members = clean_dataframe_nan(df_members)

        if not df_members.empty:
            df_members.insert(0, "STT", range(1, len(df_members) + 1))
            col_cfg_m = {"STT": st.column_config.NumberColumn("STT", width=50, disabled=True), "Họ và tên": st.column_config.TextColumn("Họ và tên", width=180, disabled=True), "Chức vụ": st.column_config.TextColumn("Chức vụ", width=90, disabled=not is_admin), "Phân môn chính": st.column_config.TextColumn("Phân môn chính", width=160, disabled=not is_admin), "Email": st.column_config.TextColumn("Email", width=180, disabled=not is_admin), "Số điện thoại": st.column_config.TextColumn("Số điện thoại", width=110, disabled=not is_admin), "Ghi chú": st.column_config.TextColumn("Ghi chú", width=90, disabled=not is_admin)}
            with st.form("form_realtime_m", border=False):
                edited_m = st.data_editor(df_members, use_container_width=True, hide_index=True, column_config=col_cfg_m, key="st_m_editor")
                if st.form_submit_button("💾 XÁC NHẬN LƯU THAY ĐỔI DANH SÁCH GIÁO VIÊN", type="primary", use_container_width=True, disabled=not is_admin):
                    conn = sqlite3.connect(DB_PATH)
                    for _, r in edited_m.iterrows():
                        conn.execute("INSERT OR REPLACE INTO org_members (fullname, position, main_subject, email, phone, note) VALUES (?, ?, ?, ?, ?, ?)", (str(r["Họ và tên"]), str(r["Chức vụ"]), str(r["Phân môn chính"]), str(r["Email"]), str(r["Số điện thoại"]), str(r["Ghi chú"])))
                    conn.commit(); conn.close(); st.success("🎉 Đã lưu!"); st.rerun()
            out_m = io.BytesIO()
            with pd.ExcelWriter(out_m, engine='openpyxl') as w: edited_m.to_excel(w, index=False, sheet_name="Danh_Sach_GV")
            st.download_button(label="📥 Tải file Excel danh sách giáo viên về máy", data=out_m.getvalue(), file_name="Danh_Sach_Thanh_Vien_To.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    with tab2:
        st.markdown("#### 📊 Sơ đồ Phân công giảng dạy của Tổ")
        excel_a = st.file_uploader("📥 Tải file Excel phân công giảng dạy lên hệ thống:", type=["xlsx", "xls"], key="up_a")
        if excel_a and st.button("🚀 Xác nhận nạp dữ liệu phân công giảng dạy", type="secondary", use_container_width=True):
            try:
                df_up_a = pd.read_excel(excel_a)
                df_up_a.columns = [str(c).strip() for c in df_up_a.columns]
                conn = sqlite3.connect(DB_PATH)
                for _, r in df_up_a.iterrows():
                    conn.execute("INSERT OR REPLACE INTO org_assignments (fullname, subject_class, homeroom, concurrent, total_periods) VALUES (?, ?, ?, ?, ?)", (str(r.get("Họ tên GV", "")), str(r.get("Môn-Lớp", "-")), str(r.get("Chủ nhiệm", "-")), str(r.get("Kiêm nhiệm", "-")), str(r.get("Tổng số tiết", "0"))))
                conn.commit(); conn.close(); st.success("🎉 Đã nạp!"); st.rerun()
            except Exception as e: st.error(f"Lỗi: {e}")

        conn = sqlite3.connect(DB_PATH)
        df_assign = pd.read_sql_query("SELECT m.fullname as [Họ tên GV], ifnull(a.subject_class, '-') as [Môn-Lớp], ifnull(a.homeroom, '-') as [Chủ nhiệm], ifnull(a.concurrent, '-') as [Kiêm nhiệm], ifnull(a.total_periods, '0') as [Tổng số tiết] FROM org_members m LEFT JOIN org_assignments a ON m.fullname = a.fullname", conn)
        conn.close()
        df_assign = clean_dataframe_nan(df_assign)
        
        if not df_assign.empty:
            df_assign.insert(0, "STT", range(1, len(df_assign) + 1))
            col_cfg_a = {"STT": st.column_config.NumberColumn("STT", width=50, disabled=True), "Họ tên GV": st.column_config.TextColumn("Họ tên GV", width=180, disabled=True), "Môn-Lớp": st.column_config.TextColumn("Môn-Lớp", width=220, disabled=not is_admin), "Chủ nhiệm": st.column_config.TextColumn("Chủ nhiệm", width=100, disabled=not is_admin), "Kiêm nhiệm": st.column_config.TextColumn("Kiêm nhiệm", width=110, disabled=not is_admin), "Tổng số tiết": st.column_config.TextColumn("Tổng số tiết", width=90, disabled=not is_admin)}
            with st.form("form_realtime_a", border=False):
                edited_a = st.data_editor(df_assign, use_container_width=True, hide_index=True, column_config=col_cfg_a, key="st_a_editor")
                if st.form_submit_button("💾 XÁC NHẬN LƯU THAY ĐỔI SƠ ĐỒ PHÂN CÔNG DẠY", type="primary", use_container_width=True, disabled=not is_admin):
                    conn = sqlite3.connect(DB_PATH)
                    for _, r in edited_a.iterrows():
                        conn.execute("INSERT OR REPLACE INTO org_assignments (fullname, subject_class, homeroom, concurrent, total_periods) VALUES (?, ?, ?, ?, ?)", (str(r["Họ tên GV"]), str(r["Môn-Lớp"]), str(r["Chủ nhiệm"]), str(r["Kiêm nhiệm"]), str(r["Tổng số tiết"])))
                    conn.commit(); conn.close(); st.success("🎉 Đã cập nhật!"); st.rerun()
            out_a = io.BytesIO()
            with pd.ExcelWriter(out_a, engine='openpyxl') as w: edited_a.to_excel(w, index=False, sheet_name="Phan_Cong")
            st.download_button(label="📥 Tải file Excel phân công giảng dạy về máy", data=out_a.getvalue(), file_name="So_Do_Phan_Cong_Giang_Day.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    with tab3:
        st.markdown("#### 🏆 Sổ theo dõi Thành tích & Thi đua qua các năm học")
        years_options = [f"{y} - {y+1}" for y in range(2020, 2036)]
        selected_year = st.selectbox("📅 CHỌN NĂM HỌC TRA CỨU THI ĐUA:", years_options, index=5)
        excel_e = st.file_uploader(f"📥 Tải file Excel thi đua năm học {selected_year} lên hệ thống:", type=["xlsx", "xls"], key="up_e")
        if excel_e and st.button(f"🚀 Xác nhận nạp dữ liệu thi đua {selected_year}", type="secondary", use_container_width=True):
            try:
                df_up_e = pd.read_excel(excel_e)
                df_up_e.columns = [str(c).strip() for c in df_up_e.columns]
                conn = sqlite3.connect(DB_PATH)
                for _, r in df_up_e.iterrows():
                    conn.execute("""
                    INSERT OR REPLACE INTO org_emulation_years (nam_hoc, fullname, dg_vien_chuc, bd_hsg, nckh, stem, sang_kien, gvdg, gvcng, tdtt_nv, hien_mau, khac)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (selected_year, str(r.get("Họ và tên", "")), str(r.get("Đánh giá viên chức", "")), str(r.get("BD HSG", "")), str(r.get("NCKH", "")), str(r.get("STEM", "")), str(r.get("Sáng kiến", "")), str(r.get("GVDG", "")), str(r.get("GVCNG", "")), str(r.get("TDTT-NV", "")), str(r.get("Hiển máu nhân đạo", "")), str(r.get("Khác", ""))))
                conn.commit(); conn.close(); st.success("🎉 Đã nạp!"); st.rerun()
            except Exception as e: st.error(f"Lỗi: {e}")

        conn = sqlite3.connect(DB_PATH)
        df_emulation = pd.read_sql_query("SELECT fullname as [Họ và tên], dg_vien_chuc as [Đánh giá viên chức], bd_hsg as [BD HSG], nckh as [NCKH], stem as [STEM], sang_kien as [Sáng kiến], gvdg as [GVDG], gvcng as [GVCNG], tdtt_nv as [TDTT-NV], hien_mau as [Hiển máu nhân đạo], khac as [Khác] FROM org_emulation_years WHERE nam_hoc = ?", conn, params=[selected_year])
        conn.close()
        df_emulation = clean_dataframe_nan(df_emulation)

        if not df_emulation.empty:
            df_emulation.insert(0, "Năm học", selected_year)
            df_emulation.insert(0, "STT", range(1, len(df_emulation) + 1))
            col_cfg_e = {"STT": st.column_config.NumberColumn("STT", width=50, disabled=True), "Năm học": st.column_config.TextColumn("Năm học", width=90, disabled=True), "Họ và tên": st.column_config.TextColumn("Họ và tên", width=180, disabled=True), "Đánh giá viên chức": st.column_config.TextColumn("Đánh giá viên chức", disabled=not is_admin), "BD HSG": st.column_config.TextColumn("BD HSG", disabled=not is_admin), "NCKH": st.column_config.TextColumn("NCKH", disabled=not is_admin), "STEM": st.column_config.TextColumn("STEM", disabled=not is_admin), "Sáng kiến": st.column_config.TextColumn("Sáng kiến", disabled=not is_admin), "GVDG": st.column_config.TextColumn("GVDG", disabled=not is_admin), "GVCNG": st.column_config.TextColumn("GVCNG", disabled=not is_admin), "TDTT-NV": st.column_config.TextColumn("TDTT-NV", disabled=not is_admin), "Hiển máu nhân đạo": st.column_config.TextColumn("Hiển máu nhân đạo", disabled=not is_admin), "Khác": st.column_config.TextColumn("Khác", disabled=not is_admin)}
            with st.form("form_realtime_e", border=False):
                edited_e = st.data_editor(df_emulation, use_container_width=True, hide_index=True, column_config=col_cfg_e, key="st_e_editor")
                if st.form_submit_button("💾 XÁC NHẬN CẬP NHẬT TOÀN BỘ SỔ THI ĐUA NĂM HỌC", type="primary", use_container_width=True, disabled=not is_admin):
                    conn = sqlite3.connect(DB_PATH)
                    for _, r in edited_e.iterrows():
                        conn.execute("""
                        INSERT OR REPLACE INTO org_emulation_years (nam_hoc, fullname, dg_vien_chuc, bd_hsg, nckh, stem, sang_kien, gvdg, gvcng, tdtt_nv, hien_mau, khac)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """, (selected_year, str(r["Họ và tên"]), str(r["Đánh giá viên chức"]), str(r["BD HSG"]), str(r["NCKH"]), str(r["STEM"]), str(r["Sáng kiến"]), str(r["GVDG"]), str(r["GVCNG"]), str(r["TDTT-NV"]), str(r["Hiển máu nhân đạo"]), str(r["Khác"])))
                    conn.commit(); conn.close(); st.success("🎉 Đã lưu!"); st.rerun()
            out_e = io.BytesIO()
            with pd.ExcelWriter(out_e, engine='openpyxl') as w: edited_e.to_excel(w, index=False, sheet_name="Thi_Dua")
            st.download_button(label=f"📥 Kết xuất Báo cáo Thi đua năm học {selected_year} (.xlsx)", data=out_e.getvalue(), file_name=f"Bao_Cao_Thi_Dua_Nam_Hoc_{selected_year.replace(' ', '')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
        else: st.info(f"ℹ️ Chưa có dữ liệu lưu trữ thi đua cho **{selected_year}**.")

# --- CÁC HÀM VỆ TINH ---
# ==================================================================================
# --- PHÂN HỆ: BIÊN BẢN SINH HOẠT TỔ CHUYÊN MÔN NÂNG CAO ---
# ==================================================================================
def setup_minutes_database():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    # Tạo bảng lưu trữ biên bản cuộc họp vĩnh viễn
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS org_minutes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        meeting_date TEXT,
        session_number TEXT UNIQUE,
        present_members TEXT,
        absent_members TEXT,
        content_text TEXT,
        resolution TEXT
    )
    """)
    conn.commit()
    conn.close()

def export_minutes_to_docx(row_data):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document()
    # Định dạng lề trang chuẩn hành chính Việt Nam (Top: 2cm, Bottom: 2cm, Left: 3cm, Right: 1.5cm)
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
        
    # Tiêu ngữ quốc hiệu căn giữa
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h1 = p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n")
    r_h1.bold = True
    r_h1.font.name = 'Times New Roman'
    r_h1.font.size = Pt(13)
    
    # Tiêu đề biên bản chữ đỏ căn giữa
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(f"\nBIÊN BẢN SINH HOẠT TỔ CHUYÊN MÔN\n(Số: {row_data[2]})\n")
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(255, 0, 0)
    
    # Nội dung biên bản căn đều 2 bên, cỡ chữ 14
    items = [
        f"- Thời gian tiến hành họp: {row_data[1]}",
        f"- Thành phần tham dự: {row_data[3]}",
        f"- Thành viên vắng mặt (lý do): {row_data[4]}",
        "\n--- NỘI DUNG DIỄN BIẾN CUỘC HỌP ---",
        f"{row_data[5]}",
        "\n--- NGHỊ QUYẾT / KẾT LUẬN CHUNG ---",
        f"{row_data[6]}"
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        if "---" in item:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
def render_meeting_minutes():
    setup_minutes_database()
    st.markdown("<h3 style='text-align: left; color: #1E3A8A;'>📝 TRÌNH LẬP BIÊN BẢN HỌC TỔ CHUYÊN MÔN NÂNG CAO</h3>", unsafe_allow_html=True)
    
    tab_lap, tab_kho = st.tabs(["✍️ Lập biên bản mới", "🗄️ Kho lưu trữ biên bản họp"])
    
    # Khởi tạo bộ nhớ tạm Session State để lưu văn bản AI sinh ra
    if "ai_resolution_draft" not in st.session_state:
        st.session_state["ai_resolution_draft"] = ""

    # ==================== THẺ 1: LẬP BIÊN BẢN MỚI ====================
    with tab_lap:
        with st.form("form_create_minutes"):
            c1, c2 = st.columns(2)
            m_date = c1.date_input("Ngày tiến hành họp tổ:")
            m_num = c2.text_input("Số biên bản (Ký hiệu):", placeholder="Ví dụ: BB-01/KHTN")
            
            m_present = st.text_area("Thành phần tham dự (Tên các thầy cô có mặt):", placeholder="Ví dụ: Thầy Dưỡng (Chủ trì), Cô Trang (Thư ký), Cô Nhi, Thầy Cường...")
            m_absent = st.text_input("Thành viên vắng mặt (Ghi rõ lý do):", placeholder="Ví dụ: Thầy Hạnh (Ốm có phép)")
            
            st.markdown("**💬 Diễn biến cuộc họp và Ý kiến đóng góp của các thành viên:**")
            m_content = st.text_area("Nội dung chi tiết cuộc họp:", placeholder="Ví dụ: 1. Thầy Dưỡng triển khai kế hoạch chuyên môn tuần 6; 2. Cô Trang báo cáo tiến độ vào điểm SMAS; 3. Thảo luận về phương pháp tích hợp năng lực số vào bài dạy...", height=200)
            
            # Khu vực gọi Trợ lý AI trợ giúp thư ký cô đọng nghị quyết
            st.markdown("🤖 **Hỗ trợ từ Trợ lý AI:**")
            col_ai_btn, _ = st.columns([1.5, 2.5])
            with col_ai_btn:
                # Nút gọi prompt chạy ngầm không làm lag trang
                run_ai_btn = st.form_submit_button("🧠 Trợ lý AI: Tóm tắt Nghị quyết cuộc họp", type="secondary")
                
            if run_ai_btn:
                if not m_content:
                    st.warning("⚠️ Vui lòng nhập nội dung chi tiết cuộc họp để AI làm căn cứ tóm tắt!")
                else:
                    with st.spinner("AI đang nghiên cứu diễn biến cuộc họp để cô đọng kết luận quyết nghị..."):
                        # Gọi hàm chạy prompt giấu khóa bảo mật
                        prompt_minutes = f"Bạn là thư ký hành chính cấp cao. Hãy đọc nội dung diễn biến cuộc họp tổ chuyên môn sau đây và tóm tắt lại thành một đoạn văn ngắn gọn, xúc tích ghi rõ các Quyết nghị/Kết luận bắt buộc tổ phải thực hiện: {m_content}. Lưu ý bỏ dấu sao sao **, ghi dòng danh sách gạch ngang."
                        try:
                            # Tận dụng hàm gọi AI dùng chung từ app.py truyền qua lambda ngầm
                            from app import run_ai_prompt_safe
                            api_key_system = st.secrets.get("GEMINI_API_KEY", "")
                            res_ai, _ = run_ai_prompt_safe(prompt_minutes, api_key_system)
                            st.session_state["ai_resolution_draft"] = res_ai
                        except:
                            st.session_state["ai_resolution_draft"] = "- Triển khai toàn diện kế hoạch dạy học tuần mới.\n- Hoàn thành hồ sơ SMAS đúng hạn quy định."

            m_resolution = st.text_area("Nghị quyết / Kết luận cuộc họp (Thầy có thể sửa lại ý AI):", value=st.session_state["ai_resolution_draft"], height=120)
            
            st.markdown("<br>", unsafe_allow_html=True)
            save_minutes = st.form_submit_button("💾 KHÓA & LƯU BIÊN BẢN VÀO HỆ THỐNG", type="primary", use_container_width=True)
            
        if save_minutes:
            if not m_num or not m_content:
                st.warning("⚠️ Vui lòng nhập đầy đủ Số biên bản và Nội dung cuộc họp!")
            else:
                conn = sqlite3.connect(DB_PATH)
                cursor = conn.cursor()
                try:
                    cursor.execute("""
                    INSERT OR REPLACE INTO org_minutes (meeting_date, session_number, present_members, absent_members, content_text, resolution)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """, (str(m_date), m_num.strip(), m_present.strip(), m_absent.strip(), m_content.strip(), m_resolution.strip()))
                    conn.commit()
                    st.success(f"✅ Biên bản số **{m_num}** đã được khóa sổ và lưu vĩnh viễn vào hệ thống!")
                    st.session_state["ai_resolution_draft"] = "" # Reset bộ đệm
                    st.rerun()
                except Exception as e:
                    st.error(f"Lỗi lưu trữ: {e}")
                finally:
                    conn.close()
    # ==================== THÈ 2: KHO LƯU TRỮ BIÊN BẢN ====================
    with tab_kho:
        st.markdown("#### 🗄️ Thư viện lưu trữ biên bản sinh hoạt tổ")
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT id, meeting_date, session_number, present_members, absent_members, content_text, resolution FROM org_minutes ORDER BY meeting_date DESC")
        all_minutes = cursor.fetchall()
        conn.close()
        
        if not all_minutes:
            st.info("Hiện tại chưa có biên bản cuộc họp nào được lưu trữ trong phiên này.")
        else:
            for idx, row in enumerate(all_minutes):
                col_exp, col_del = st.columns([0.88, 0.12])
                with col_exp:
                    with st.expander(f"📝 {idx+1}. Biên bản số: {row[2]} - Ngày họp: {row[1]}"):
                        st.markdown(f"**🗓️ Ngày họp:** {row[1]}")
                        st.markdown(f"**👥 Thành phần tham gia:** {row[3]}")
                        st.markdown(f"**❌ Vắng mặt:** {row[4]}")
                        st.markdown("**📄 Diễn biến cuộc họp:**")
                        st.caption(row[5])
                        st.markdown("**🏅 Quyết nghị cuộc họp:**")
                        st.write(row[6])
                        
                        # Tích hợp xuất file Word chuẩn hành chính 
                        from org_manager import export_minutes_to_docx
                        word_data = export_minutes_to_docx(row)
                        st.download_button(
                            label="📥 Tải file Word (.docx) biên bản này",
                            data=word_data,
                            file_name=f"Bien_Ban_So_{row[2].replace('/', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_word_min_{row[0]}"
                        )
                with col_del:
                    st.write("") # Cân dòng
                    if st.button("🗑️ Xóa", key=f"del_min_{row[0]}", use_container_width=True):
                        conn = sqlite3.connect(DB_PATH)
                        conn.execute("DELETE FROM org_minutes WHERE id=?", (row[0],))
                        conn.commit()
                        conn.close()
                        st.success("Đã xóa biên bản!")
                        st.rerun()

# --- HÀM GIỮ NGUYÊN GỌI TỪ APP.PY ---
def render_personal_plan():
    st.header("📅 KẾ HOẠCH GIÁO DỤC CÁ NHÂN")

def render_personal_plan():
    st.header("📅 KẾ HOẠCH GIÁO DỤC CÁ NHÂN")
