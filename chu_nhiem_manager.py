import streamlit as st
import io
import re
from docx import Document

# ==========================================
# PHẦN 1: KẾ HOẠCH NĂM HỌC
# ==========================================
def render_nam_hoc_tab():
    st.markdown("### 📌 I. ĐẶC ĐIỂM TÌNH HÌNH LỚP")
    col_tl, col_kk = st.columns(2)
    with col_tl:
        st.text_area(
            "Thuận lợi", 
            value="- Nhà trường luôn quan tâm đến công tác chủ nhiệm, có những kế hoạch chỉ đạo kịp thời đến tập thể lớp.\n- Tập thể lớp 9D có tinh thần đoàn kết, có ý thức trong công việc được giao, công việc của lớp, có tinh thần cầu tiến, có ý thức vươn lên trong học tập và rèn luyện đạo đức tác phong.\n- Có đội ngũ ban cán sự lớp nhiệt tình năng nổ, hết lòng vì công việc chung.",
            height=150, key="nam_hoc_thuan_loi"
        )
    with col_kk:
        st.text_area(
            "Khó khăn", 
            value="Một số học sinh vẫn còn tâm lý lười học, thiếu tập trung, đặc biệt ở các môn khoa học.\nCó học sinh thuộc diện hoàn cảnh khó khăn, thiếu sự quan tâm của gia đình.\nMột số em có biểu hiện chống đối nhẹ, dễ bị lôi kéo.\nKỹ năng tự học, tự quản và làm việc nhóm của học sinh chưa đồng đều.\nLớp có một số em ý thức chưa được tốt.",
            height=150, key="nam_hoc_kho_khan"
        )
        
    st.write("---")
    st.markdown("### 📝 II. KẾ HOẠCH NĂM HỌC 2025-2026")
    
    col_rl, col_mdc = st.columns(2)
    with col_rl:
        st.text_area("Rèn luyện trong nhà trường", value="+ Xây dựng uy tín với học sinh, với đồng nghiệp, với cha mẹ học sinh và xã hội về chuyên môn, nghiệp vụ. Tư cách đạo đức, tác phong sư phạm mẫu mực trong sinh hoạt.", height=120, key="nam_hoc_ren_luyen")
    with col_mdc:
        st.text_area("Mục đích yêu cầu chung", value="- Luôn kính trọng người trên, thầy cô giáo, cán bộ và nhân viên nhà trường; thương yêu và giúp đỡ nhau, có ý thức xây dựng tập thể, đoàn kết với các bạn, được các bạn tin yêu.", height=120, key="nam_hoc_muc_dich_chung")
        
    col_ctc, col_bpc = st.columns(2)
    with col_ctc:
        st.text_area("Chỉ tiêu chung", value="100% học sinh không vi phạm nội quy trường học.\n95% học sinh đạt Rèn luyện tốt và Khá.\nKhông có học sinh bỏ học.", height=100, key="nam_hoc_chi_tieu_chung")
    with col_bpc:
        st.text_area("Biện pháp chính chung", value="- GVCN quán triệt cho học sinh nội quy của nhà trường, bài học văn hóa ứng xử, giáo dục truyền thống nhà trường, giáo dục môi trường, an toàn giao thông.", height=100, key="nam_hoc_bien_phap_chung")

    st.markdown("#### 🎯 Mục đích cụ thể từng mặt")
    st.info("✨ Phân hệ 1: HỌC TẬP")
    col_ht_yc, col_ht_ct = st.columns(2)
    with col_ht_yc:
        st.text_area("Yêu cầu (Học tập)", value="- 100% học sinh đi học có đầy đủ SGK và các dụng cụ phục vụ học tập.\n- 100% soạn bài và làm bài tập, học thuộc bài ở nhà.\n- 100% HS tích cực tham gia phát biểu xây dựng bài trong các tiết học.", height=120, key="ht_yeu_cau")
    with col_ht_ct:
        st.text_area("Chỉ tiêu (Học tập)", value="- Mức Tốt: 5 HS\n- Mức Khá: 12 HS\n- Mức Đạt: 22 HS trở lên\n- Danh hiệu 'Học sinh Xuất sắc': 1 HS", height=120, key="ht_chi_tieu")
    st.text_area("Biện pháp chính (Học tập)", value="- Tổ chức cho HS thảo luận, tìm phương pháp học tập đạt hiệu quả cao.\n- Kết hợp với GVBM thành lập nhóm hs cán sự theo bộ môn, đẩy mạnh các phong trào 'Đôi bạn học tập'.", height=100, key="ht_bien_phap")

    st.write("")
    st.info("✨ Phân hệ 2: GIÁO DỤC HƯỚNG NGHIỆP")
    col_hn_yc, col_hn_ct = st.columns(2)
    with col_hn_yc:
        st.text_area("Yêu cầu (Hướng nghiệp)", value="- Giáo dục học sinh có ý thức yêu lao động, có trách nhiệm giữ gìn vệ sinh, bảo vệ tài sản, cảnh quan, môi trường thiên nhiên.\n- Giúp học sinh định hướng được nghề nghiệp phù hợp.", height=120, key="hn_yeu_cau")
    with col_hn_ct:
        st.text_area("Chỉ tiêu (Hướng nghiệp)", value="- 100% HS làm tốt công tác bảo vệ tài sản, cảnh quan, môi trường 'Xanh - sạch - đẹp'.\n- 100% HS tham gia đầy đủ các buổi lao động.", height=120, key="hn_chi_tieu")
    st.text_area("Biện pháp chính (Hướng nghiệp)", value="- Quán triệt học sinh tham gia đầy đủ các buổi lao động trường, lớp, ý thức bảo vệ môi trường, lao động tập thể.\n- Nhận xét, đánh giá, khen thưởng cho HS khi tham gia.", height=100, key="hn_bien_phap")

    st.write("")
    st.info("✨ Phân hệ 3: GIÁO DỤC NGOÀI GIỜ LÊN LỚP")
    col_gl_yc, col_gl_ct = st.columns(2)
    with col_gl_yc:
        st.text_area("Yêu cầu (Layout Ngoài giờ lên lớp)", value="- Kĩ năng tự nhận thức, giải quyết tình huống khó khăn trong cuộc sống.\n- Kĩ năng giao tiếp và thương thuyết (bao hàm tính tự kiềm chế).\n- Kĩ năng lựa chọn và quyết định.", height=120, key="gl_yeu_cau")
    with col_gl_ct:
        st.text_area("Chỉ tiêu (Ngoài giờ lên lớp)", value="100% HS rèn kĩ năng sống cần được giáo dục và rèn luyện.", height=120, key="gl_chi_tieu")
    st.text_area("Biện pháp chính (Ngoài giờ lên lớp)", value="- Nâng cao ý thức tự nguyện, tự giác, tự chủ, phát huy tính tích cực trong mọi hoạt động.\n- Không chỉ rèn luyện cho mình mà quan tâm đến việc rèn luyện chung của cả tập thể.", height=100, key="gl_bien_phap")

    st.write("---")
    st.markdown("### 📊 III. CHỈ TIÊU TOÀN DIỆN CUỐI NĂM")
    col_td_ct, col_td_bp = st.columns(2)
    with col_td_ct:
        st.text_area("Chỉ tiêu cuối năm", value="1. Danh hiệu lớp: Lớp Tiên tiến - Chi đội: Mạnh\n2. Rèn luyện: Tốt: 40 (97,7%) Khá: 02 (2,3%)\n3. Kết quả học tập: Mức Tốt: 5HS (13,5%) Khá: 12 HS (27,9%)", height=120, key="td_chi_tieu")
    with col_td_bp:
        st.text_area("Biện pháp chính cuối năm", value="- GVCN thường xuyên bám sát lớp, gần gũi tiếp xúc với HS để tìm hiểu hoàn cảnh.\n- Phối hợp chặt chẽ với GVBM và Tổng phụ trách Đội.\n- Phối hợp thường xuyên với PHHS.", height=120, key="td_bien_phap")
        
    if st.button("💾 Lưu Toàn bộ Kế hoạch năm học", type="primary", key="btn_save_nam_hoc_full"):
        st.success("🎉 Đã lưu trữ thành công Kế hoạch năm học!")

# ==========================================
# PHẦN 2: KẾ HOẠCH THÁNG (CÓ AI)
# ==========================================
def export_to_word(title, content_text):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in str(content_text).split("\n"):
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def render_thang_tab(run_ai_prompt_safe=None):
    st.write("#### 🛠 CẤU HÌNH THÔNG TIN CHỦ NHIỆM")
    col_khoi, col_lop, col_thang = st.columns(3)
    
    with col_khoi:
        selected_khoi = st.selectbox("Chọn Khối lớp:", ["Khối lớp 6", "Khối lớp 7", "Khối lớp 8", "Khối lớp 9"], key="sb_khoi_lop")
    with col_lop:
        lop_dict = {
            "Khối lớp 6": ["6A","6B","6C","6D","6E","6F"], 
            "Khối lớp 7": ["7A","7B","7C","7D","7E","7F"], 
            "Khối lớp 8": ["8A","8B","8C","8D","8E","8F"], 
            "Khối lớp 9": ["9A","9B","9C","9D","9E","9F","9G"]
        }
        selected_lop = st.selectbox("Chọn Lớp chủ nhiệm:", lop_dict.get(selected_khoi, ["6A"]), key="sb_lop_chu_nhiem")
    with col_thang:
        thang_options = [f"Tháng {i}/2026" for i in range(8, 13)] + [f"Tháng {i}/2027" for i in range(1, 6)]
        selected_thang = st.selectbox("Chọn Tháng công tác:", thang_options, key="sb_thang_cong_tac")
        
    st.write("---")
    ghi_chu_them = st.text_input("Yêu cầu bổ sung đặc biệt cho tháng này (nếu có):", placeholder="Ví dụ: Tập trung nề nếp thi đua...", key="txt_ghi_chu_them")
    
    if "ta_main_editor" not in st.session_state:
        st.session_state["ta_main_editor"] = ""
        
    if st.button("🚀 Khởi tạo Kế hoạch bằng AI", type="primary", key="btn_chu_nhiem_ai"):
        if run_ai_prompt_safe is not None:
            with st.spinner(f"AI đang thiết lập kế hoạch {selected_thang}..."):
                # Cập nhật Prompt ép cấu trúc Dấu đầu dòng lùi lề
                prompt_he_thong = f"""
                Bạn là trợ lý AI cho giáo viên chủ nhiệm THCS Việt Nam. Hãy lập bản kế hoạch công tác chủ nhiệm chi tiết cho lớp {selected_lop} trong {selected_thang}.
                
                LƯU Ý QUAN TRỌNG VỀ THUẬT NGỮ VÀ ĐỊNH DẠNG (BẮT BUỘC TUÂN THỦ):
                1. TUYỆT ĐỐI KHÔNG dùng từ "Học lực", phải dùng từ "Kết quả học tập".
                2. TUYỆT ĐỐI KHÔNG dùng từ "Hạnh kiểm", phải dùng từ "Rèn luyện".
                3. KHÔNG dùng ký tự ** để in đậm. 
                
                YÊU CẦU ĐẦU RA THEO ĐÚNG CẤU TRÚC VÀ DẤU ĐẦU DÒNG SAU (Copy y hệt format này):
                KẾ HOẠCH THÁNG {selected_thang.replace('Tháng ', '')}
                1. Chủ điểm: [Tên chủ điểm giáo dục tương ứng tháng]
                [Nhiệm vụ thi đua]
                
                2. Nội dung hoạt động:
                [Đầu việc lớn]
                
                KẾ HOẠCH TỪNG TUẦN:
                TUẦN 1: (Từ ngày/tháng - ngày/tháng)
                * Nề nếp, sĩ số:
                   - [Nội dung chi tiết 1]
                   - [Nội dung chi tiết 2]
                * Hoạt động cụ thể:
                   - [Nội dung chi tiết 1]
                   - [Nội dung chi tiết 2]
                
                Ghi chú từ GV: {ghi_chu_them}
                """
                
                raw_response = run_ai_prompt_safe(prompt_he_thong)
                
                if raw_response is None:
                    final_text = "Lỗi: Không nhận được phản hồi từ AI."
                else:
                    if isinstance(raw_response, tuple):
                        final_text = str(raw_response[0])
                    elif hasattr(raw_response, 'text'):
                        final_text = str(raw_response.text)
                    else:
                        final_text = str(raw_response)
                        
                    if final_text.startswith("('"):
                        final_text = final_text[2:]
                    if final_text.endswith("',)"):
                        final_text = final_text[:-3]
                    elif final_text.endswith("')"):
                        final_text = final_text[:-2]
                        
                    final_text = final_text.replace("\\n", "\n")
                    
                    # BỘ LỌC HẬU KỲ XỬ LÝ TEXT:
                    # 1. Quét và thay thế các từ cấm
                    final_text = final_text.replace("Học lực", "Kết quả học tập").replace("học lực", "kết quả học tập")
                    final_text = final_text.replace("Hạnh kiểm", "Rèn luyện").replace("hạnh kiểm", "rèn luyện")
                    
                    # 2. Xóa tất cả các dấu sao in đậm
                    final_text = final_text.replace("**", "")
                    
                    # 3. Ép định dạng các đầu mục chính:
                    # Chữ TUẦN: Xóa mọi dấu - hoặc * ở đầu
                    final_text = re.sub(r'^[\-\*\s]+TUẦN', 'TUẦN', final_text, flags=re.MULTILINE)
                    # Chữ Nề nếp và Hoạt động: Ép phải có * ở đầu
                    final_text = re.sub(r'^[\-\*\s]*Nề nếp', '* Nề nếp', final_text, flags=re.MULTILINE)
                    final_text = re.sub(r'^[\-\*\s]*Hoạt động', '* Hoạt động', final_text, flags=re.MULTILINE)
                    
                st.session_state["ta_main_editor"] = final_text.strip()
        else:
            st.info("Hệ thống kết nối AI đang được đồng bộ...")

    st.write("#### 📝 KHUNG VĂN BẢN KẾ HOẠCH THÁNG (Xem trước & Sửa đổi)")
    
    edited_text = st.text_area(
        label="Nội dung kế hoạch công tác (Nhấn vào để sửa đổi):",
        height=450,
        key="ta_main_editor"
    )
    
    if isinstance(st.session_state["ta_main_editor"], str) and st.session_state["ta_main_editor"].strip():
        st.write("")
        file_name_doc = f"Ke_hoach_chu_nhiem_{selected_lop}_{selected_thang.replace('/', '_')}.docx"
        word_file_bytes = export_to_word(f"KẾ HOẠCH CHỦ NHIỆM LỚP {selected_lop} - {selected_thang.upper()}", edited_text)
        
        st.download_button(
            label="📥 Tải xuống file Word (.docx)",
            data=word_file_bytes,
            file_name=file_name_doc,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="btn_download_word"
        )

# ==========================================
# PHẦN 3: HÀM GỌI CHÍNH TỪ APP.PY
# ==========================================
def render_chu_nhiem_section(run_ai_prompt_safe=None):
    tab_tong_quan, tab_hang_thang = st.tabs([
        "📊 Đặc điểm tình hình & Kế hoạch năm học", 
        "📅 Kế hoạch công tác theo Tháng (AI Tự động)"
    ])
    
    with tab_tong_quan:
        render_nam_hoc_tab()

    with tab_hang_thang:
        render_thang_tab(run_ai_prompt_safe)
