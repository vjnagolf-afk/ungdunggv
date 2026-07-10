import streamlit as st
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re

def set_paragraph_spacing(paragraph, before_pt=3.0, after_pt=4.5):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    p_pr.append(spacing)

def export_rubric_to_docx(title_text, markdown_content):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)

    MAU_DO = RGBColor(255, 0, 0)
    MAU_XANH_DUONG = RGBColor(0, 51, 153)
    MAU_DEN = RGBColor(0, 0, 0)

    # Tiêu đề Rubric căn giữa màu đỏ
    p_title = doc.add_paragraph()
    set_paragraph_spacing(p_title, 6.0, 6.0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title_text.upper())
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = MAU_DO

    lines = markdown_content.split('\n')
    in_table = False
    table_data = []

    for line in lines:
        clean_line = line.strip().replace('**', '').replace('###', '').replace('##', '').replace('#', '')
        
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if '---|' in line or ':---|' in line: continue
            in_table = True
            cells = [c.strip().replace('**', '') for c in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            if in_table and table_data:
                num_rows = len(table_data)
                num_cols = len(table_data) if num_rows > 0 else 0
                if num_cols > 0:
                    word_table = doc.add_table(rows=num_rows, cols=num_cols)
                    word_table.style = 'Table Grid'
                    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for r_idx, row in enumerate(table_data):
                        for c_idx, val in enumerate(row):
                            if c_idx < num_cols:
                                cell = word_table.cell(r_idx, c_idx)
                                cell.text = val
                                for para in cell.paragraphs:
                                    set_paragraph_spacing(para, 2.0, 3.0)
                                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    for r in para.runs:
                                        r.font.name = 'Times New Roman'
                                        r.font.size = Pt(14)
                                        r.font.color.rgb = MAU_DEN
                in_table = False
                table_data = []

        if not clean_line: continue

        if re.match(r'^(I|II|III|IV|V)\.', clean_line) or re.match(r'^\d+\.', clean_line):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 4.0, 4.5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean_line)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = MAU_XANH_DUONG
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(clean_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
def render_assessment_section(run_ai_prompt_safe_func):
    st.markdown("<h3 style='text-align: center; color: red;'>🎯 TRỢ LÝ THIẾT KẾ RUBRIC ĐÁNH GIÁ HỌC SINH TỰ ĐỘNG</h3>", unsafe_allow_html=True)
    
    tab_thiet_ke, tab_thu_vien = st.tabs(["📝 THIẾT KẾ TIÊU CHÍ RUBRIC AI", "🗄️ LƯU TRỮ RUBRIC"])
    
    if "ket_qua_rubric" not in st.session_state: st.session_state["ket_qua_rubric"] = ""
    if "lich_su_rubric" not in st.session_state: st.session_state["lich_su_rubric"] = []

    with tab_thiet_ke:
        st.write("Nhập tên nội dung bài học hoặc chủ đề để hệ thống tự động thiết kế bảng tiêu chí đánh giá định lượng.")
        
        noi_dung = st.text_input("Tên nội dung kiến thức bài học / Chương / Chủ đề / Sản phẩm:", placeholder="Ví dụ: Mô hình xe phản lực - Chương Tốc độ chuyển động")
        
        col_lop, col_loai = st.columns(2)
        with col_lop:
            lop = st.text_input("Lớp:", placeholder="Ví dụ: 7A")
        with col_loai:
            hinh_thuc = st.selectbox("Hình thức đánh giá:", ["Qua sản phẩm học tập", "Qua bài thuyết trình", "Qua hoạt động nhóm", "Đánh giá năng lực thực hành"])

        col_btn1, col_blank, col_btn2 = st.columns([2.0, 1.3, 1.7])
        
        with col_btn2:
            st.write(""); st.write("")
            nut_rubric = st.button("⚡ Khởi tạo tiêu chí bằng AI", type="primary", use_container_width=True)

        if nut_rubric:
            if not noi_dung:
                st.warning("⚠️ Vui lòng điền nội dung kiến thức bài học hoặc tên sản phẩm!")
            else:
                with st.spinner("🧠 Hệ thống đang phân tích mục tiêu bài học để lập bảng Rubric..."):
                    prompt_rubric = f"""
                    Bạn là Chuyên gia kiểm tra đánh giá giáo dục phổ thông tại Việt Nam. Hãy thiết kế một bảng tiêu chí đánh giá định lượng (Rubric) chi tiết cho:
                    - Nội dung/Sản phẩm: {noi_dung}
                    - Lớp: {lop}
                    - Hình thức đánh giá: {hinh_thuc}
                    
                    TIÊU ĐỀ KHỐI (Viết in hoa ở dòng đầu):
                    BẢNG TIÊU CHÍ RUBRIC ĐÁNH GIÁ: {noi_dung.upper()} - LỚP {lop.upper()}
                    
                    YÊU CẦU CẤU TRÚC VÀ ĐỊNH DẠNG VĂN BẢN:
                    1. BẢNG TIÊU CHÍ RUBRIC: Phải xây dựng dưới dạng bảng Markdown hoàn chỉnh bằng ký tự '|'. Các cột bao gồm: Tiêu chí đánh giá, Trọng số (%), Mức Tốt (8.0-10 điểm), Mức Khá (6.5-7.9 điểm), Mức Đạt (5.0-6.4 điểm), Mức Chưa đạt (Dưới 5.0 điểm).
                    2. NỘI DUNG TIÊU CHÍ: Phải bám sát đặc thù kiến thức bài học/sản phẩm của học sinh lớp {lop}. Mô tả rõ ràng hành vi định lượng để giáo viên dễ chấm, không viết chung chung.
                    3. BỎ TOÀN BỘ các ký tự dấu sao kép '**' ở đầu và cuối các từ hoặc tiêu đề cột để tránh lỗi font chữ đậm.
                    4. Danh sách liệt kê dùng dấu gạch ngang '-' ở đầu dòng.
                    """
                    ket_qua, model = run_ai_prompt_safe_func(prompt_rubric)
                    st.session_state["ket_qua_rubric"] = ket_qua

        with col_btn1:
            st.write(""); st.write("")
            from danh_gia_manager import export_rubric_to_docx
            title_file = f"Rubric_{lop}_{noi_dung[:20].replace(' ', '_')}" if noi_dung else "Rubric_Danh_Gia"
            docx_data = export_rubric_to_docx(f"BẢNG RUBRIC ĐÁNH GIÁ: {noi_dung}", st.session_state["ket_qua_rubric"]) if st.session_state["ket_qua_rubric"] else b""
            st.download_button(
                label="📥 Tải file Word (.docx) chuẩn về máy",
                data=docx_data,
                file_name=f"{title_file}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=(st.session_state["ket_qua_rubric"] == ""),
                use_container_width=True
            )

        st.markdown("**📊 Bảng tiêu chí và thang điểm chi tiết sinh bởi AI:**")
        with st.container(border=True):
            if st.session_state["ket_qua_rubric"]:
                st.markdown(st.session_state["ket_qua_rubric"])
                if st.button("📥 Lưu vào thư viện hệ thống", use_container_width=True):
                    st.session_state["lich_su_rubric"].append({"Nội dung": noi_dung, "Lớp": lop, "Data": st.session_state["ket_qua_rubric"]})
                    st.success("✅ Đã lưu bảng Rubric thành công!")
            else:
                st.caption("Nội dung bảng Rubric sẽ hiển thị tại đây sau khi khởi tạo bằng AI...")

    with tab_thu_vien:
        st.markdown("### 🗄️ DANH SÁCH CÁC TIÊU CHÍ RUBRIC ĐÃ LƯU TRỮ")
        if not st.session_state["lich_su_rubric"]:
            st.info("Chưa có tiêu chí nào được lưu trong phiên làm việc này.")
        else:
            for idx, item in enumerate(st.session_state["lich_su_rubric"]):
                col_exp, col_del = st.columns([0.88, 0.12])
                with col_exp:
                    with st.expander(f"📊 {idx+1}. {item['Nội dung']} - Lớp {item['Lớp']}"):
                        st.markdown(item["Data"])
                        from danh_gia_manager import export_rubric_to_docx
                        saved_docx = export_rubric_to_docx(f"BẢNG RUBRIC ĐÁNH GIÁ: {item['Nội dung']}", item["Data"])
                        st.download_button(
                            label="📥 Tải lại file Word (.docx)",
                            data=saved_docx,
                            file_name=f"Luu_tru_Rubric_{idx+1}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_saved_rubric_{idx}"
                        )
                with col_del:
                    st.write("")
                    if st.button("🗑️ Xóa", key=f"del_rubric_{idx}", use_container_width=True):
                        st.session_state["lich_su_rubric"].pop(idx)
                        st.rerun()
