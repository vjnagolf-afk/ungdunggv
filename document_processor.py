# document_processor.py
import re
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
from math_compiler import generate_plot_stream, process_runs_with_math
def read_uploaded_docx(uploaded_file):
    try:
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    except: return "Lỗi đọc file Word"

def read_uploaded_pdf(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except: return "Lỗi đọc file PDF"

def export_to_docx_vietnam_standard(text_content, title_name, school_name, group_name="TỔ KHOA HỌC TỰ NHIÊN"):
    text_content = re.sub(r'(?m)^#+\s*', '', text_content)
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14) # 🚀 CẬP NHẬT CỠ CHỮ CHUẨN ĐỒNG BỘ 14PT
    
    admin_table = doc.add_table(rows=1, cols=2)
    admin_table.autofit = False
    
    cell_l = admin_table.rows[0].cells
    p_left = cell_l[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.add_run(f"{school_name.upper()}\n").bold = True
    p_left.add_run(f"{group_name.upper()}\n").bold = True
    p_left.add_run("Số: ..... /BB-TCM").font.size = Pt(11)
    
    cell_r = admin_table.rows[0].cells
    p_right = cell_r[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").bold = True
    p_right.add_run("Độc lập - Tự do - Hạnh phúc\n").bold = True
    p_right.add_run("***************").font.size = Pt(11)
    
    admin_table.columns[0].width = Inches(3.2)
    admin_table.columns[1].width = Inches(3.8)
    
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run(title_name.upper()).bold = True
    p_title.runs[0].font.size = Pt(14)
    
    in_table = False
    table_data = []
    
    def build_table():
        if not table_data: return
        cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_val in enumerate(row):
                if c_idx < cols:
                    cell = table.cell(r_idx, c_idx)
                    p = cell.paragraphs[0]
                    p.text = ""
                    process_runs_with_math(p, cell_val.strip())
        doc.add_paragraph()
        
    for line in text_content.split('\n'):
        cleaned_line = line.strip()
        if not cleaned_line: 
            continue
            
        if cleaned_line.startswith('|') and cleaned_line.endswith('|'):
            in_table = True
            row_data = [cell.strip() for cell in cleaned_line.split('|')[1:-1]]
            if all(re.match(r'^[-: ]+$', cell) for cell in row_data): continue
            table_data.append(row_data)
            continue
        
        if in_table:
            build_table()
            in_table = False
            table_data = []
            
        if '[GRAPH:' in cleaned_line:
            match = re.search(r'\[GRAPH:\s*(.+?)\]', cleaned_line)
            if match:
                eq = match.group(1)
                img_stream = generate_plot_stream(eq)
                doc.add_picture(img_stream, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if re.search(r'A\..*B\..*C\..*D\.', cleaned_line):
            sub_choices = re.split(r'(?=A\.|B\.|C\.|D\.)', cleaned_line)
            for choice in sub_choices:
                if choice.strip():
                    p = doc.add_paragraph()
                    process_runs_with_math(p, choice.strip())
            continue
            
        p = doc.add_paragraph()
        if re.match(r'^(Câu \d+:)', cleaned_line) or re.match(r'^(Câu \d+ \([^)]+\):)', cleaned_line):
            match_prefix = re.match(r'^(Câu \d+\s*(?:\([^)]+\))?:)', cleaned_line)
            prefix = match_prefix.group(1)
            rest = cleaned_line[len(prefix):]
            run_p = p.add_run(prefix + " ")
            run_p.bold = True
            process_runs_with_math(p, rest.strip())
        else:
            process_runs_with_math(p, cleaned_line)
            
    if in_table and table_data:
        build_table()
        
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
