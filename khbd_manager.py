# khbd_manager.py - ĐOẠN 1: CẤU HÌNH & TRÍCH XUẤT TÀI LIỆU (BẢN VÁ TOÀN DIỆN)
import streamlit as st
import docx  
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
from pypdf import PdfReader
import gspread
from datetime import datetime

# Nhúng bộ biên dịch toán và đồ thị thông minh để dùng cho bài soạn
from math_compiler import process_runs_with_math, generate_plot_stream

# ================= CẤU HÌNH GOOGLE SHEETS =================
SHEET_ID = '1C6642jk_oQ0g9UC2By2qsNxxfQVR0MrZYj52tRdWDlY' 

def get_khbd_sheet():
    try:
        creds_dict = None
        priority_keys = ["gspread_credentials", "GSPREAD_CREDENTIALS", "google_sheet_creds", "google_creds", "GOOGLE_KEY"]
        for key in priority_keys:
            if key in st.secrets:
                creds_dict = st.secrets[key]
                break
                
        if creds_dict is None:
            for key in st.secrets.keys():
                node = st.secrets[key]
                if hasattr(node, "get") or isinstance(node, dict):
                    if node.get("type") == "service_account" or "private_key" in node:
                        creds_dict = node
                        break
                        
        if creds_dict is None:
            return None
            
        gc = gspread.service_account_from_dict(creds_dict)
        sh = gc.open_by_key(SHEET_ID)
        return sh.worksheet("KHBD")
    except:
        return None

# --- HÀM TRÍCH XUẤT VĂN BẢN VÀ LỌC ẢNH TRÙNG LẶP ---
def extract_context_from_uploaded_files(uploaded_files):
    combined_text = ""
    extracted_images = [] 
    seen_image_sizes = set()
    for file in uploaded_files:
        try:
            if file.name.endswith('.docx'):
                doc = docx.Document(file)
                for paragraph in doc.paragraphs:
                    if paragraph.text: combined_text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        text_row = [cell.text for cell in row.cells]
                        combined_text += " | ".join(text_row) + "\n"
                for rel in doc.part.relations.values():
                    if "image" in rel.target_ref:
                        img_blob = rel.target_part.blob
                        if len(img_blob) not in seen_image_sizes:
                            seen_image_sizes.add(len(img_blob))
                            extracted_images.append(img_blob)
            elif file.name.endswith('.pdf'):
                reader = PdfReader(file)
                for page in reader.pages:
                    combined_text += (page.extract_text() or "") + "\n"
                    for img_file_object in page.images:
                        img_blob = img_file_object.data
                        if len(img_blob) not in seen_image_sizes:
                            seen_image_sizes.add(len(img_blob))
                            extracted_images.append(img_blob)
            elif file.name.endswith('.txt'):
                combined_text += file.read().decode("utf-8") + "\n"
        except Exception as e:
            st.error(f"Lỗi khi xử lý file {file.name}: {str(e)}")
    return combined_text, extracted_images

def set_paragraph_spacing(paragraph, before_pt=3.0, after_pt=4.5):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    p_pr.append(spacing)
# khbd_manager.py - ĐOẠN 2: THUẬT TOÁN KẾT XUẤT GIÁO ÁN CHUẨN WORD (VÁ LỖI ĐỀ MỤC & TOÁN HỌC)
# khbd_manager.py - ĐOẠN 2: THUẬT TOÁN KẾT XUẤT GIÁO ÁN CHUẨN ĐỊNH DẠNG MÀU SẮC YÊU CẦU
# khbd_manager.py - ĐOẠN 2: THUẬT TOÁN KẾT XUẤT GIÁO ÁN CHUẨN ĐỊNH DẠNG MÀU SẮC YÊU CẦU
# khbd_manager.py - ĐOẠN 2: THUẬT TOÁN KẾT XUẤT GIÁO ÁN VĂN BẢN XUÔI (TUYỆT ĐỐI KHÔNG DÙNG BẢNG)
# khbd_manager.py - ĐOẠN 2: THUẬT TOÁN KẾT XUẤT GIÁO ÁN CHUẨN ĐỊNH DẠNG MÀU SẮC VÀ TOÁN HỌC DỨT ĐIỂM
# khbd_manager.py - ĐOẠN 2: THUẬT TOÁN KẾT XUẤT GIÁO ÁN VĂN BẢN XUÔI SẠCH LỖI BIẾN
# khbd_manager.py - ĐOẠN 2: THUẬT TOÁN KẾT XUẤT VĂN BẢN XUÔI THEO KHUNG CỨNG TUYỆT ĐỐI CÔNG VĂN 5512
def export_khbd_to_docx(markdown_content, images_list):
    markdown_content = re.sub(r'(?m)^#+\s*', '', markdown_content)
    markdown_content = markdown_content.replace("<br>", "\n").replace("<br/>", "\n")
    
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)

    MAU_DO = RGBColor(255, 0, 0)
    MAU_XANH_DUONG = RGBColor(0, 51, 153)
    MAU_DEN = RGBColor(0, 0, 0)

    lines = markdown_content.split('\n')
    processed_lines = []
    
    # BỘ LỌC TƯỚC DẤU VÀ HỦY ĐỊNH DẠNG AI THỪA
    for line in lines:
        # Hủy toàn bộ markdown in đậm của AI để ép theo template cứng
        cl = line.strip().replace('**', '').replace('*', '') 
        if not cl or all(c in '-| ' for c in cl): 
            continue
        
        match_bullet = re.match(r'^[-–—*•+]\s*(.*)', cl)
        if match_bullet:
            core_text = match_bullet.group(1).strip()
            # Danh sách CẤM có dấu gạch ngang đầu dòng (I. II. a) b) Bước 1...)
            if re.match(r'^(I|II|III|IV|V|VI|VII)\.', core_text) or \
               re.match(r'^\d+\.\s+(Kiến|Năng|Phẩm)', core_text, re.IGNORECASE) or \
               re.match(r'^[a-d]\)\s*', core_text, re.IGNORECASE) or \
               re.match(r'^\d+\.\s+HOẠT', core_text, re.IGNORECASEChào bạn, mình hoàn toàn hiểu sự bực bội của bạn khi file xuất ra vẫn chưa đạt mức hoàn hảo. Sau khi rà soát lại **toàn bộ** hệ thống dựa trên phản hồi và các hình ảnh bạn cung cấp, mình đã tìm ra 2 nguyên nhân cốt lõi khiến hệ thống liên tục "trật nhịp":

1. **Lỗi phá vỡ dấu Toán học từ file Manager:** Trong file `khbd_manager.py` cũ, có một dòng lệnh Regex cố gắng làm sạch khoảng trắng quanh dấu `$` nhưng lại vô tình "bẻ gãy" luôn các thẻ `$$...$$` và `\[...\]`. Điều này khiến `math_compiler.py` không nhận diện được công thức để dịch sang Equation.
2. **Lỗi tràn định dạng Word:** Hàm lùi lề và bôi đậm trước đó dùng điều kiện nhận diện quá lỏng lẻo, khiến Word không khóa cứng được lề (Left: 0, Right: 0) và tô màu sai các đoạn văn bản đi kèm phía sau.

Để khắc phục dứt điểm cho **tất cả các module hiện tại và sau này**, bạn cần thay thế toàn bộ mã của cả 2 file: `khbd_manager.py` và `math_compiler.py`. Hệ thống mới được thiết kế theo cơ chế **"Khuôn đúc cứng"**, bất chấp AI sinh ra văn bản lộn xộn hay khoảng trắng thừa, code vẫn sẽ nắn lại đúng chuẩn.

---

### 1. FILE `khbd_manager.py` (Khuôn đúc định dạng cứng)
Bản này đã **xóa bỏ hoàn toàn** các lệnh can thiệp Toán học sai lầm, khóa chặt lề `Left: 0, Right: 0` và áp dụng hệ thống Rule ưu tiên (chỉ bôi đậm/tô màu đúng tiền tố, giữ nguyên hậu tố).

Bạn hãy chép đè toàn bộ nội dung file `khbd_manager.py`:

```python
import streamlit as st
import docx  
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
from pypdf import PdfReader
import gspread
from datetime import datetime

from math_compiler import process_runs_with_math, generate_plot_stream

# ================= CẤU HÌNH GOOGLE SHEETS =================
SHEET_ID = '1C6642jk_oQ0g9UC2By2qsNxxfQVR0MrZYj52tRdWDlY' 

def get_khbd_sheet():
    try:
        creds_dict = None
        priority_keys = ["gspread_credentials", "GSPREAD_CREDENTIALS", "google_sheet_creds", "google_creds", "GOOGLE_KEY"]
        for key in priority_keys:
            if key in st.secrets:
                creds_dict = st.secrets[key]
                break
                
        if creds_dict is None:
            for key in st.secrets.keys():
                node = st.secrets[key]
                if hasattr(node, "get") or isinstance(node, dict):
                    if node.get("type") == "service_account" or "private_key" in node:
                        creds_dict = node
                        break
                        
        if creds_dict is None: return None
        gc = gspread.service_account_from_dict(creds_dict)
        sh = gc.open_by_key(SHEET_ID)
        return sh.worksheet("KHBD")
    except:
        return None

# --- HÀM TRÍCH XUẤT VĂN BẢN VÀ LỌC ẢNH TRÙNG LẶP ---
def extract_context_from_uploaded_files(uploaded_files):
    combined_text = ""
    extracted_images = [] 
    seen_image_sizes = set()
    for file in uploaded_files:
        try:
            if file.name.endswith('.docx'):
                doc = docx.Document(file)
                for paragraph in doc.paragraphs:
                    if paragraph.text: combined_text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        text_row = [cell.text for cell in row.cells]
                        combined_text += " | ".join(text_row) + "\n"
                for rel in doc.part.relations.values():
                    if "image" in rel.target_ref:
                        img_blob = rel.target_part.blob
                        if len(img_blob) not in seen_image_sizes:
                            seen_image_sizes.add(len(img_blob))
                            extracted_images.append(img_blob)
            elif file.name.endswith('.pdf'):
                reader = PdfReader(file)
                for page in reader.pages:
                    combined_text += (page.extract_text() or "") + "\n"
                    for img_file_object in page.images:
                        img_blob = img_file_object.data
                        if len(img_blob) not in seen_image_sizes:
                            seen_image_sizes.add(len(img_blob))
                            extracted_images.append(img_blob)
            elif file.name.endswith('.txt'):
                combined_text += file.read().decode("utf-8") + "\n"
        except Exception as e:
            st.error(f"Lỗi khi xử lý file {file.name}: {str(e)}")
    return combined_text, extracted_images

def set_paragraph_spacing(paragraph, before_pt=3.0, after_pt=4.5):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    p_pr.append(spacing)

# ================= ĐỘNG CƠ KẾT XUẤT ĐỊNH DẠNG CỨNG =================
def export_khbd_to_docx(markdown_content, images_list):
    markdown_content = re.sub(r'(?m)^#+\s*', '', markdown_content)
    markdown_content = markdown_content.replace("<br>", "\n").replace("<br/>", "\n")
    
    # TUYỆT ĐỐI KHÔNG DÙNG REGEX CAN THIỆP DẤU $ Ở ĐÂY NỮA. CHUYỂN GIAO CHO MATH_COMPILER.

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)

    MAU_DO = RGBColor(255, 0, 0)
    MAU_XANH_DUONG = RGBColor(0, 51, 153)
    MAU_DEN = RGBColor(0, 0, 0)

    lines = markdown_content.split('\n')
    
    for line in lines:
        cl = line.strip().replace('**', '').replace('*', '') 
        if not cl or all(c in '-| ' for c in cl): 
            continue

        # 1. TƯỚC BỎ KÝ TỰ GẠCH NGANG / BULLET CHO CÁC MỤC CỐ ĐỊNH
        bullet_match = re.match(r'^[-–—*•+]\s*(.*)', cl)
        if bullet_match:
            core = bullet_match.group(1).strip()
            # Nếu là đề mục chính, xóa luôn dấu gạch
            if re.match(r'^(I|II|III|IV|V|VI|VII)\.|^\d+\.\s+(Kiến|Năng|Phẩm)|^[a-d]\)\s*|^\d+\.\s+HOẠT|^\d+\.\d+\.?\s+Hoạt|^Bước|^\d+\.\s+Đối', core, re.IGNORECASE):
                cl = core
            else:
                cl = "- " + core # Chuẩn hóa gạch ngang cho nội dung thường

        p = doc.add_paragraph()
        set_paragraph_spacing(p, 3.0, 4.5)
        # KHÓA CHẶT LỀ 0 TUYỆT ĐỐI
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.right_indent = Inches(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if '[GRAPH:' in cl:
            match = re.search(r'\[GRAPH:\s*(.+?)\]', cl)
            if match:
                img_stream = generate_plot_stream(match.group(1))
                doc.add_picture(img_stream, width=Inches(4.5))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        upper_cl = cl.upper()
        
        # RULE 1: KẾ HOẠCH BÀI DẠY, BÀI 8 (Đỏ, Đậm, Giữa)
        if upper_cl.startswith("KẾ HOẠCH BÀI DẠY") or upper_cl.startswith("BÀI"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(upper_cl)
            run.bold = True
            run.font.name = 'Times New Roman'; run.font.size = Pt(14)
            run.font.color.rgb = MAU_DO
            continue

        # RULE 2: MÔN HỌC, LỚP, THỜI LƯỢNG, TIẾT (Xanh dương, Đậm, Giữa)
        if any(upper_cl.startswith(x) for x in ["MÔN HỌC", "LỚP", "THỜI LƯỢNG", "TIẾT"]):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(upper_cl)
            run.bold = True
            run.font.name = 'Times New Roman'; run.font.size = Pt(14)
            run.font.color.rgb = MAU_XANH_DUONG
            continue

        # RULE 3: I. MỤC TIÊU BÀI HỌC, II. THIẾT BỊ... (Xanh dương, Đậm, Hoa)
        if re.match(r'^(I|II|III|IV|V|VI|VII)\.', cl):
            process_runs_with_math(p, upper_cl)
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'; r.font.size = Pt(14)
                r.font.color.rgb = MAU_XANH_DUONG
            continue

        # RULE 4: 1. Kiến thức, 2. Năng lực... (Tiền tố: Đỏ Đậm)
        m = re.match(r'^(\d+\.\s+(?:Kiến thức|Năng lực|Phẩm chất)[^:]*:?)(.*)', cl, re.IGNORECASE)
        if m:
            r_pref = p.add_run(m.group(1))
            r_pref.bold = True
            r_pref.font.name = 'Times New Roman'; r_pref.font.size = Pt(14)
            r_pref.font.color.rgb = MAU_DO
            if m.group(2).strip(): process_runs_with_math(p, m.group(2))
            continue

        # RULE 5: 1. HOẠT ĐỘNG 1... (Tiền tố: Xanh dương, Đậm, Hoa)
        m = re.match(r'^(\d+\.\s+HOẠT ĐỘNG\s+\d+:\s*[^.(0-9]*)(.*)', cl, re.IGNORECASE)
        if m:
            r_pref = p.add_run(m.group(1).upper())
            r_pref.bold = True
            r_pref.font.name = 'Times New Roman'; r_pref.font.size = Pt(14)
            r_pref.font.color.rgb = MAU_XANH_DUONG
            if m.group(2).strip(): process_runs_with_math(p, m.group(2))
            continue

        # RULE 6: 2.1 Hoạt động 1. (Tiền tố: Xanh dương, Đậm, Thường)
        m = re.match(r'^(\d+\.\d+\.?\s+Hoạt động\s+\d+\.?)(.*)', cl, re.IGNORECASE)
        if m:
            r_pref = p.add_run(m.group(1) + " ")
            r_pref.bold = True
            r_pref.font.name = 'Times New Roman'; r_pref.font.size = Pt(14)
            r_pref.font.color.rgb = MAU_XANH_DUONG
            if m.group(2).strip(): process_runs_with_math(p, m.group(2).strip())
            continue

        # RULE 7: Tiền tố Đen Đậm (a) Mục tiêu, Bước 1, 1. Đối với...)
        m = re.match(r'^([a-d]\)\s*(?:Năng lực|Mục tiêu|Nội dung|Sản phẩm|Tổ chức thực hiện)[^:]*:?|Bước\s+\d+:?|\d+\.\s+Đối với[^:]*:?)(.*)', cl, re.IGNORECASE)
        if m:
            pref = m.group(1)
            if not pref.endswith(' '): pref += ' '
            r_pref = p.add_run(pref)
            r_pref.bold = True
            r_pref.font.name = 'Times New Roman'; r_pref.font.size = Pt(14)
            r_pref.font.color.rgb = MAU_DEN
            if m.group(2).strip(): process_runs_with_math(p, m.group(2).strip())
            continue

        # RULE 8: Nội dung thông thường
        process_runs_with_math(p, cl)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            if not r.font.size: r.font.size = Pt(14)
            
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def render_khbd_section(run_ai_prompt_safe_func):
    st.markdown("<h3 style='text-align: center; color: blue;'>🧠 TRỢ LÝ THIẾT KẾ KẾ HOẠCH BÀI DẠY (KHBD) AI PHÁT TRIỂN NĂNG LỰC</h3>", unsafe_allow_html=True)
    
    tab_thiet_ke, tab_thu_vien = st.tabs(["📝 THIẾT KẾ KHBD TỰ ĐỘNG", "🗄️ THƯ VIỆN BÀI SOẠN"])
    
    if "ket_qua_khbd" not in st.session_state: st.session_state["ket_qua_khbd"] = ""
    if "lich_su_khbd" not in st.session_state: st.session_state["lich_su_khbd"] = []
    if "images_khbd" not in st.session_state: st.session_state["images_khbd"] = []

    with tab_thiet_ke:
        st.write("Nhập thông tin bài học và tải lên tài liệu tham khảo để AI lập tiến trình dạy học Công văn 5512.")
        
        ten_bai = st.text_input("Tên bài học / Chủ đề bài dạy:", placeholder="Ví dụ: Bài 4: Tốc độ chuyển động - Khoa học tự nhiên 7")
        
        col_lop, col_bo = st.columns(2)
        with col_lop:
            lop_khbd = st.text_input("Lớp dạy:", value="Lớp 7")
        with col_bo:
            bo_sach = st.selectbox("Bộ sách giáo khoa:", ["Kết nối tri thức với cuộc sống", "Cánh Diều", "Chân trời sáng tạo", "Chương trình GDPT 2018"])

        col_tg, col_loai = st.columns(2)
        with col_tg:
            thoi_luong = st.text_input("Thời lượng bài dạy (Tiết):", value="2 tiết")
        with col_loai:
            kieu_khbd = st.selectbox("Mẫu cấu trúc thiết kế:", ["Chuẩn Công văn 5512 (Đầy đủ 4 hoạt động)", "Rút gọn (Tiết kiệm thời gian)", "Giáo án hoạt động trải nghiệm/STEM"])

        st.markdown("**Tải lên tài liệu tham khảo thô (Đề cương, Sách, file nội dung bài học nếu có):**")
        files_tailieu = st.file_uploader("Chọn tệp (.docx, .pdf, .txt):", type=["docx", "pdf", "txt"], accept_multiple_files=True, key="khbd_files_upload")

        context_data = ""
        if files_tailieu:
            context_data, st.session_state["images_khbd"] = extract_context_from_uploaded_files(files_tailieu)
            st.success(f"📊 Đã nạp thành công văn bản tham khảo và trích xuất được {len(st.session_state['images_khbd'])} hình ảnh!")

        col_chk1, col_chk2 = st.columns(2)
        with col_chk1:
            chk_ai_digital = st.checkbox("Tích hợp năng lực số và AI", value=True)
        with col_chk2:
            chk_strict_file = st.checkbox("Bám sát 100% file tài liệu tải lên", value=False)

        yeu_cau_rieng = st.text_area("Yêu cầu sư phạm bổ sung (Tùy chọn):", placeholder="Ví dụ: Thiết kế thêm một trò chơi khởi động sôi nổi...")

        col_btn1, col_blank, col_btn2 = st.columns([2.2, 1.0, 1.8])
        
        with col_btn2:
            st.write(""); st.write("")
            nut_tao_khbd = st.button("⚡ Tiến hành thiết kế bài dạy bằng AI", type="primary", use_container_width=True)

        if nut_tao_khbd:
            if not ten_bai:
                st.warning("⚠️ Vui lòng điền tên bài học hoặc chủ đề giảng dạy!")
            else:
                with st.spinner("🧠 AI đang phân tích dữ liệu, bám sát Công văn 5512 để thiết kế tiến trình..."):
                    prompt_requirements = ""
                    if chk_ai_digital:
                        prompt_requirements += "\n- TẠI MỤC I.2 (VỀ NĂNG LỰC): Bổ sung tiểu mục 'c) Năng lực số và AI'."
                    if chk_strict_file:
                        prompt_requirements += f"\n- QUY ĐỊNH BẮT BUỘC: Bám sát 100% tệp tài liệu: {context_data[:3000]}"

                    prompt_khbd = f"""
                    Bạn là Chuyên gia Phương pháp dạy học. Hãy soạn Kế hoạch bài dạy chi tiết:
                    - Tên bài: {ten_bai}
                    - Lớp: {lop_khbd} - Bộ sách: {bo_sach}
                    - Thời lượng: {thoi_luong}
                    - Yêu cầu: {prompt_requirements}
                    - Bổ sung: {yeu_cau_rieng if yeu_cau_rieng else 'Không có'}
                    - CÔNG THỨC TOÁN HỌC: Đặt tất cả công thức/đại lượng trong $...$, $$...$$, \(...\) hoặc \[...\].
                    """
                    ket_qua, _ = run_ai_prompt_safe_func(prompt_khbd)
                    st.session_state["ket_qua_khbd"] = ket_qua

        if st.session_state["ket_qua_khbd"]:
            st.markdown("### 📋 Giáo án bài dạy xem trước:")
            st.markdown(st.session_state["ket_qua_khbd"])

        with col_btn1:
            st.write(""); st.write("")
            title_file_khbd = f"KHBD_{lop_khbd}_{ten_bai[:20].replace(' ', '_')}" if ten_bai else "Ke_Hoach_Bai_Day"
            
            if st.session_state["ket_qua_khbd"]:
                docx_data_khbd = export_khbd_to_docx(st.session_state["ket_qua_khbd"], st.session_state["images_khbd"])
                is_disabled_khbd = False
                sheet_khbd = get_khbd_sheet()
                if sheet_khbd is not None:
                    try: sheet_khbd.append_row([ten_bai, lop_khbd, bo_sach, thoi_luong, datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
                    except: pass
            else:
                docx_data_khbd = b""
                is_disabled_khbd = True

            st.download_button(
                label="📥 Tải tệp Giáo án Word (.docx)",
                data=docx_data_khbd,
                file_name=f"{title_file_khbd}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=is_disabled_khbd,
                use_container_width=True
            )

    with tab_thu_vien:
        st.markdown(f"🔗 [Mở Google Sheets quản lý bài soạn](https://docs.google.com/spreadsheets/d/{SHEET_ID}/edit)")
def render_khbd_section(run_ai_prompt_safe_func):
    st.markdown("<h3 style='text-align: center; color: blue;'>🧠 TRỢ LÝ THIẾT KẾ KẾ HOẠCH BÀI DẠY (KHBD) AI PHÁT TRIỂN NĂNG LỰC</h3>", unsafe_allow_html=True)
    
    tab_thiet_ke, tab_thu_vien = st.tabs(["📝 THIẾT KẾ KHBD TỰ ĐỘNG", "🗄️ THƯ VIỆN BÀI SOẠN"])
    
    if "ket_qua_khbd" not in st.session_state: st.session_state["ket_qua_khbd"] = ""
    if "lich_su_khbd" not in st.session_state: st.session_state["lich_su_khbd"] = []
    if "images_khbd" not in st.session_state: st.session_state["images_khbd"] = []

    with tab_thiet_ke:
        st.write("Nhập thông tin bài học và tải lên tài liệu tham khảo để AI lập tiến trình dạy học Công văn 5512.")
        
        ten_bai = st.text_input("Tên bài học / Chủ đề bài dạy:", placeholder="Ví dụ: Bài 4: Tốc độ chuyển động - Khoa học tự nhiên 7")
        
        col_lop, col_bo = st.columns(2)
        with col_lop:
            lop_khbd = st.text_input("Lớp dạy:", value="Lớp 7")
        with col_bo:
            bo_sach = st.selectbox("Bộ sách giáo khoa:", ["Kết nối tri thức với cuộc sống", "Cánh Diều", "Chân trời sáng tạo", "Chương trình GDPT 2018"])

        col_tg, col_loai = st.columns(2)
        with col_tg:
            thoi_luong = st.text_input("Thời lượng bài dạy (Tiết):", value="2 tiết")
        with col_loai:
            kieu_khbd = st.selectbox("Mẫu cấu trúc thiết kế:", ["Chuẩn Công văn 5512 (Đầy đủ 4 hoạt động)", "Rút gọn (Tiết kiệm thời gian)", "Giáo án hoạt động trải nghiệm/STEM"])

        st.markdown("**Tải lên tài liệu tham khảo thô (Đề cương, Sách, file nội dung bài học nếu có):**")
        files_tailieu = st.file_uploader("Chọn tệp (.docx, .pdf, .txt):", type=["docx", "pdf"], accept_multiple_files=True, key="khbd_files_upload")

        context_data = ""
        if files_tailieu:
            context_data, st.session_state["images_khbd"] = extract_context_from_uploaded_files(files_tailieu)
            st.success(f"📊 Đã nạp thành công văn bản tham khảo và trích xuất được {len(st.session_state['images_khbd'])} hình ảnh!")

        col_chk1, col_chk2 = st.columns(2)
        with col_chk1:
            chk_ai_digital = st.checkbox("Tích hợp năng lực số và AI", value=True)
        with col_chk2:
            chk_strict_file = st.checkbox("Bám sát 100% file tài liệu tải lên", value=False)

        yeu_cau_rieng = st.text_area("Yêu cầu sư phạm bổ sung (Tùy chọn):", placeholder="Ví dụ: Thiết kế thêm một trò chơi khởi động sôi nổi; lồng ghép công thức toán học chi tiết...")

        col_btn1, col_blank, col_btn2 = st.columns([2.2, 1.0, 1.8])
        
        with col_btn2:
            st.write(""); st.write("")
            nut_tao_khbd = st.button("⚡ Tiến hành thiết kế bài dạy bằng AI", type="primary", use_container_width=True)

        if nut_tao_khbd:
            if not ten_bai:
                st.warning("⚠️ Vui lòng điền tên bài học hoặc chủ đề giảng dạy!")
            else:
                with st.spinner("🧠 AI đang phân tích dữ liệu, bám sát Công văn 5512 để thiết kế tiến trình..."):
                    prompt_requirements = ""
                    if chk_ai_digital:
                        prompt_requirements += """
                        - TẠI MỤC I.2 (VỀ NĂNG LỰC): Bắt buộc phải bổ sung thêm riêng một tiểu mục: 'c) Năng lực số và AI' (Mô tả chi tiết tiêu chí học sinh biết khai thác tài nguyên số hoặc ứng dụng AI gì để thực hành, học tập trong bài soạn này).
                        - TẠI TIẾN TRÌNH CÁC HOẠT ĐỘNG DẠY HỌC: Thiết kế lồng ghép hoạt động giao nhiệm vụ yêu cầu học sinh sử dụng ứng dụng mô phỏng hoặc công cụ công nghệ số.
                        """
                    if chk_strict_file:
                        prompt_requirements += f"\n- QUY ĐỊNH BẮT BUỘC: Khai thác và bám sát chính xác 100% kiến thức từ tệp tài liệu tham khảo thô sau đây: {context_data[:3000]}"

                    prompt_khbd = f"""
                    Bạn là Chuyên gia Phương pháp dạy học cao cấp tại Việt Nam. Hãy soạn một Kế hoạch bài dạy (KHBD) chi tiết theo Công văn 5512 cho bài học:
                    - Tên bài: {ten_bai}
                    - Lớp: {lop_khbd} - Thuộc bộ sách: {bo_sach}
                    - Thời lượng: {thoi_luong}
                    - Kiểu cấu trúc mẫu: {kieu_khbd}
                    
                    YÊU CẦU ĐỊNH HƯỚNG SƯ PHẠM QUY ĐỊNH (THI SÁT TUYỆT ĐỐI):
                    {prompt_requirements}
                    - Yêu cầu bổ sung riêng: {yeu_cau_rieng if yeu_cau_rieng else "Không có"}
                    
                    QUY ĐỊNH ĐỊNH DẠNG CÔNG THỨC TOÁN HỌC & ĐỒ THỊ (Nghiêm ngặt):
                    - Tuyệt đối không viết chữ toán dạng phẳng dính văn bản thường. Tất cả công thức như v = s/t, các phân số chồng tầng, đại lượng vật lý bắt buộc phải đặt cô lập gọn gàng trong cặp dấu đô-la $...$ độc lập (Ví dụ: $v = \\frac{{s}}{{t}}$, $1 \\text{{ m/s}} = 3.6 \\text{{ km/h}}$).
                    - Nếu có đồ thị hàm số, xuất chuỗi dạng [GRAPH: tên_hàm] (Ví dụ: [GRAPH: 3*x]) để Word tự động vẽ hình.
                    """
                    ket_qua, model_thuc_te = run_ai_prompt_safe_func(prompt_khbd)
                    st.session_state["ket_qua_khbd"] = ket_qua

        if st.session_state["ket_qua_khbd"]:
            st.markdown("### 📋 Giáo án bài dạy xem trước:")
            st.markdown(st.session_state["ket_qua_khbd"])

        with col_btn1:
            st.write(""); st.write("")
            title_file_khbd = f"KHBD_{lop_khbd}_{ten_bai[:20].replace(' ', '_')}" if ten_bai else "Ke_Hoach_Bai_Day"
            
            if st.session_state["ket_qua_khbd"]:
                docx_data_khbd = export_khbd_to_docx(st.session_state["ket_qua_khbd"], st.session_state["images_khbd"])
                is_disabled_khbd = False
                
                sheet_khbd = get_khbd_sheet()
                if sheet_khbd is not None:
                    try:
                        sheet_khbd.append_row([ten_bai, lop_khbd, bo_sach, thoi_luong, datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
                    except: pass
            else:
                docx_data_khbd = b""
                is_disabled_khbd = True

            st.download_button(
                label="📥 Tải tệp Giáo án Word (.docx) bản chuẩn hành chính",
                data=docx_data_khbd,
                file_name=f"{title_file_khbd}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=is_disabled_khbd,
                use_container_width=True
            )

    with tab_thu_vien:
        st.write("🗄️ Quản lý kho lưu trữ giáo án trực tuyến:")
        # 🚀 ĐÃ SỬA CHÍNH XÁC ĐƯỜNG DẪN URL LIÊN KẾT TRANG TÍNH GOOGLE SHEETS TẠI ĐÂY
        st.markdown(f"🔗 [Bấm vào đây để mở trực tiếp Google Sheets quản lý bài soạn](https://docs.google.com/spreadsheets/d/{SHEET_ID}/edit)")

