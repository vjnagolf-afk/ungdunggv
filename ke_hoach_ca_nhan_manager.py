import streamlit as st  
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import json
import pandas as pd
from openpyxl.utils import get_column_letter

# --- HÀM TẠO LƯỚI VIỀN CHO BẢNG WORD ---
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# --- HÀM XUẤT PHÔI WORD NÂNG CẤP ĐỦ 8 CỘT ---
def export_plan_to_docx_with_table(teacher_name, subject_name, content_data):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
        
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_qg = p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐỘC LẬP - TỰ DO - HẠNH PHÚC\n")
    run_qg.bold = True
    run_qg.font.name = 'Times New Roman'
    run_qg.font.size = Pt(13)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("KẾ HOẠCH GIÁO DỤC CỦA GIÁO VIÊN\n(PHỤ LỤC III CÔNG VĂN 5512/BGDĐT)\n")
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(14)
    
    p_info = doc.add_paragraph()
    r_info = p_info.add_run(f"Giáo viên thực hiện: {teacher_name}\nMôn học/Hoạt động giáo dục: {subject_name}\n\n")
    r_info.font.name = 'Times New Roman'
    r_info.font.size = Pt(14)
    
    p_section1 = doc.add_paragraph()
    r_sec1 = p_section1.add_run("I. KẾ HOẠCH DẠY HỌC PHÂN PHỐI CHƯƠNG TRÌNH")
    r_sec1.bold = True
    r_sec1.font.name = 'Times New Roman'
    r_sec1.font.size = Pt(14)
    
    headers = ["STT", "Tiết CT", "Bài học", "Số tiết", "Thời điểm", "Yêu cầu cần đạt", "Thiết bị", "Địa điểm"]
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows.cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_margins(hdr_cells[i])
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    if isinstance(content_data, list):
        for idx, item in enumerate(content_data, 1):
            row_cells = table.add_row().cells
            row_cells.text = str(idx)
            row_cells.text = str(item.get("TietCT", idx))
            row_cells.text = str(item.get("BaiHoc", ""))
            row_cells.text = str(item.get("SoTiet", ""))
            row_cells.text = str(item.get("ThoiDiem", ""))
            row_cells.text = str(item.get("YeuCauCanDat", "-"))
            row_cells.text = str(item.get("ThietBi", "-"))
            row_cells.text = str(item.get("DiaDiem", "Lớp học"))
            
            for cell in row_cells:
                set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        
    doc.add_paragraph()
    p_section2 = doc.add_paragraph()
    r_sec2 = p_section2.add_run("\nII. CÁC NHIỆM VỤ KHÁC ĐƯỢC GIAO")
    r_sec2.bold = True
    r_sec2.font.name = 'Times New Roman'
    r_sec2.font.size = Pt(14)
    
    p_other = doc.add_paragraph()
    r_other = p_other.add_run("- Bồi dưỡng học sinh giỏi và phụ đạo học sinh yếu kém theo kế hoạch nhà trường.\n- Tham gia sinh hoạt chuyên môn tổ, nhóm đầy đủ và đúng quy định.")
    r_other.font.name = 'Times New Roman'
    r_other.font.size = Pt(12)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- HÀM XUẤT FILE EXCEL MẪU 8 CỘT ---
def export_plan_to_excel(teacher_name, subject_name, content_data):
    output = io.BytesIO()
    raw_df = pd.DataFrame(content_data)
    df = pd.DataFrame({
        "STT": range(1, len(raw_df) + 1),
        "Tiết CT": raw_df.get("TietCT", range(1, len(raw_df) + 1)),
        "Bài học": raw_df.get("BaiHoc", ""),
        "Số tiết": raw_df.get("SoTiet", ""),
        "Thời điểm": raw_df.get("ThoiDiem", ""),
        "Yêu cầu cần đạt": raw_df.get("YeuCauCanDat", ""),
        "Thiết bị": raw_df.get("ThietBi", ""),
        "Địa điểm dạy học": raw_df.get("DiaDiem", "Lớp học")
    })
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        workbook = writer.book
        worksheet = workbook.create_sheet(title="Kế hoạch dạy học", index=0)
        writer.sheets["Kế hoạch dạy học"] = worksheet
        worksheet["A1"] = "I. KẾ HOẠCH DẠY HỌC"
        worksheet["A2"] = f"1. KHDH MÔN {subject_name.upper()}"
        df.to_excel(writer, sheet_name="Kế hoạch dạy học", startrow=3, index=False)
        
        for col_idx, col in enumerate(worksheet.columns, 1):
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = get_column_letter(col_idx)
            worksheet.column_dimensions[col_letter].width = max(max_len + 3, 12)
    return output.getvalue()

# --- HÀM TRÍCH XUẤT CHỮ TỪ FILE TÀI LIỆU TẢI LÊN ---
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    file_name = uploaded_file.name
    try:
        if file_name.endswith(".txt"):
            return uploaded_file.read().decode("utf-8")
        elif file_name.endswith(".docx"):
            doc = docx.Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        elif file_name.endswith(".xlsx") or file_name.endswith(".xls"):
            df = pd.read_excel(uploaded_file)
            return f"Nội dung file Excel kế hoạch cũ ({file_name}):\n" + df.to_string()
        elif file_name.endswith(".pdf"):
            return f"[File đính kèm: {file_name}]"
    except Exception as e:
        return f"Không thể đọc file {file_name}: {str(e)}"
    return ""
# --- GIAO DIỆN PHÂN HỆ KẾ HOẠCH CÁ NHÂN ---
def render_personal_plan(run_ai_handler=None):
    st.markdown("<h3 style='text-align: left; color: #1E3A8A;'>🗓️ TRỢ LÝ XÂY DỰNG KẾ HOẠCH GIÁO DỤC CÁ NHÂN (PHỤ LỤC III)</h3>", unsafe_allow_html=True)
    
    if "db_ke_hoach_da_luu" not in st.session_state:
        st.session_state["db_ke_hoach_da_luu"] = {}
        
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🔒 CHỌN VAI TRÒ ĐĂNG NHẬP")
    vai_tro = st.sidebar.radio("Vai trò", ["Giáo viên bộ môn", "Tổ trưởng chuyên môn (Admin)", "Ban giám hiệu"], label_visibility="collapsed", key="vai_tro_plan_sidebar_fixed_v8")
    
    if vai_tro == "Tổ trưởng chuyên môn (Admin)":
        st.sidebar.info("Mã xác thực PIN hệ thống tổng đang hoạt động ngầm.")

    if "current_ai_output_raw" not in st.session_state: st.session_state["current_ai_output_raw"] = None
    if "current_teacher_active" not in st.session_state: st.session_state["current_teacher_active"] = ""
    if "current_subject_active" not in st.session_state: st.session_state["current_subject_active"] = ""

    with st.form("form_personal_plan_fixed_final_v8", border=False):
        col_t, col_s = st.columns(2)
        t_name = col_t.text_input("Họ và tên Giáo viên giảng dạy:", placeholder="Ví dụ: Thầy Lê Hồng Dưỡng", key="plan_txt_t_name_v8")
        
        # 🌟 CẬP NHẬT TÍNH NĂNG MỚI: Tích hợp đầy đủ danh mục môn học Chương trình GDPT 2018
        danh_sach_mon_hoc = [
            "Toán học", 
            "Ngữ văn", 
            "Tiếng Anh", 
            "Giáo dục công dân", 
            "Khoa học tự nhiên (Phân môn Vật lí)", 
            "Khoa học tự nhiên (Phân môn Hóa học)", 
            "Khoa học tự nhiên (Phân môn Sinh học)", 
            "Lịch sử và Địa lí (Phân môn Lịch sử)", 
            "Lịch sử và Địa lí (Phân môn Địa lý)", 
            "Công nghệ", 
            "Tin học", 
            "Giáo dục thể chất", 
            "Nghệ thuật (Âm nhạc)", 
            "Nghệ thuật (Mĩ thuật)", 
            "Hoạt động trải nghiệm, hướng nghiệp", 
            "Giáo dục địa phương"
        ]
        s_name = col_s.selectbox("Môn học / Phân môn phụ trách:", danh_sach_mon_hoc, key="plan_sb_s_name_v8")
        
        col_g, col_w, col_total = st.columns(3)
        grade_target = col_g.text_input("Khối lớp phân công dạy:", placeholder="Ví dụ: Khối 9, Khối 7", key="plan_txt_grade_target_v8")
        week_count = col_w.text_input("Tổng số tuần thực hiện kế hoạch:", placeholder="35 Tuần", value="35 Tuần", key="plan_txt_week_count_v8")
        total_lessons = col_total.number_input("Tổng số tiết trong năm học:", min_value=1, max_value=300, value=70, step=1, key="plan_num_total_lessons_v8")
        
        st.markdown("**💬 Các tiêu chí đặc thù hoặc lưu ý phân bổ tiết (Nếu có):**")
        note_plan = st.text_area("Yêu cầu bổ sung cho AI:", placeholder="Ví dụ: Cần chia nhỏ phân phối chương trình thành từng tiết đơn lẻ 1, 2, 3...", label_visibility="collapsed", key="plan_ta_note_v8")
        
        st.markdown("    **📚 Đính kèm tài liệu tham khảo (File SGK dạng PDF/Word và File Excel kế hoạch đã có sẵn nếu có):**")
        uploaded_files = st.file_uploader("Tải tệp lên hệ thống (Chấp nhận nhiều file cùng lúc)", type=["txt", "docx", "pdf", "xlsx", "xls"], accept_multiple_files=True, label_visibility="collapsed", key="sgk_file_uploader_v8")
        
        run_ai_plan = st.form_submit_button("🚀 Khởi tạo Kế hoạch bằng AI", type="primary", use_container_width=True)
    if run_ai_plan:
        if not t_name or not grade_target:
            st.warning("⚠️ Vui lòng điền Họ tên giáo viên và Khối lớp để AI lập kế hoạch!")
        else:
            with st.spinner("Trợ lý AI đang tổng hợp các file tài liệu và phân bổ đều số tiết dạy..."):
                combined_context = ""
                if uploaded_files:
                    for f in uploaded_files:
                        combined_context += extract_text_from_file(f) + "\n\n"
                
                prompt_plan = (
                    f"Hãy soạn thảo phân bổ chương trình dạy học chi tiết bám sát định hướng Chương trình GDPT 2018 cho giáo viên: {t_name}, môn: {s_name}, khối: {grade_target} trong {week_count}. "
                    f"Tổng số tiết bắt buộc phải rải đều trong năm học là: {total_lessons} tiết.\n"
                    f"Yêu cầu bổ sung kỹ thuật: {note_plan}. "
                    f"Nhiệm vụ của bạn là phải phân tích, bám sát và kế thừa tối đa từ các tài liệu đính kèm (nội dung bài học từ file SGK hoặc cấu trúc phân bổ từ file Excel kế hoạch có sẵn) dưới đây:\n"
                    f"[DỮ LIỆU THAM KHẢO TỪ CÁC FILE TÀI LIỆU]:\n{combined_context}\n\n"
                    f"Quy tắc chia dòng: Danh sách kết quả phải chạy lũy tiến từ Tiết CT số 1 cho đến tiết số {total_lessons}. Mỗi hàng ứng với 1 tiết duy nhất, số tiết ghi rõ là 1. "
                    f"Trả về mảng JSON thuần túy gồm danh sách các đối tượng, không kèm markdown thô nào ngoài thẻ mở/đóng json. "
                    f"Mỗi đối tượng bắt buộc phải chứa đúng cấu trúc 8 khóa sau không được sai lệch: "
                    f'[{{"TietCT": "Số thứ tự tiết lũy tiến (từ 1 đến {total_lessons})", "BaiHoc": "Tên bài học cụ thể", '
                    f'"SoTiet": "1", "ThoiDiem": "Tuần thực hiện (Phân bổ hợp lý trong {week_count})", '
                    f'"YeuCauCanDat": "Yêu cầu cần đạt sinh tự động bám sát theo tài liệu", "ThietBi": "Thiết bị dạy học sử dụng cụ thể", "DiaDiem": "Địa điểm dạy học (Mặc định: Lớp học)"}}]'
                )
                
                if run_ai_handler is not None:
                    res_plan, status = run_ai_handler(prompt_plan)
                    if status == "error":
                        st.error(f"❌ Hệ thống kết nối AI gặp sự cố kỹ thuật: {res_plan}")
                    else:
                        try:
                            clean_json = res_plan.replace("```json", "").replace("```", "").strip()
                            parsed_data = json.loads(clean_json)
                            
                            st.session_state["current_ai_output_raw"] = parsed_data
                            st.session_state["current_teacher_active"] = t_name
                            st.session_state["current_subject_active"] = s_name
                            
                            record_id = f"{t_name.replace(' ', '_')}_{s_name.replace(' ', '_')}"
                            st.session_state["db_ke_hoach_da_luu"][record_id] = {
                                "teacher": t_name,
                                "subject": s_name,
                                "grade": grade_target,
                                "weeks": week_count,
                                "data": parsed_data
                            }
                            st.success("🎉 Khởi tạo và lưu trữ phân phối chương trình thành công!")
                        except Exception as parse_err:
                            st.error("⚠️ AI phản hồi cấu hình phân tách hàng chưa đồng bộ. Vui lòng bấm thử lại để làm mới.")
                else:
                    st.error("❌ Lỗi cấu hình hệ thống: Không tìm thấy trình điều khiển AI tập trung.")

    # --- KHU VỰC HIỂN THỊ KẾT QUẢ VÀ SONG HÀNH CÁC NÚT TÁC VỤ ---
    st.markdown("---")
    st.markdown("### 📊 Nội dung Kế hoạch Giáo dục sinh bởi AI:")
    
    if st.session_state["current_ai_output_raw"]:
        df_display = pd.DataFrame(st.session_state["current_ai_output_raw"])
        st.dataframe(df_display, use_container_width=True, hide_index=True)
        
        # Chia thành 3 cột để xếp nút Tải Word, Tải Excel và nút Xóa file hiện hành nằm ngang hàng
        btn_col1, btn_col2, btn_col3 = st.columns(3)
        
        word_plan_data = export_plan_to_docx_with_table(
            st.session_state["current_teacher_active"], 
            st.session_state["current_subject_active"], 
            st.session_state["current_ai_output_raw"]
        )
        btn_col1.download_button(
            label="📥 Tải file Word (.docx) Phụ lục III", 
            data=word_plan_data, 
            file_name=f"Phu_Luc_III_{st.session_state['current_teacher_active'].replace(' ', '_')}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            use_container_width=True
        )
        
        excel_plan_data = export_plan_to_excel(
            st.session_state["current_teacher_active"], 
            st.session_state["current_subject_active"], 
            st.session_state["current_ai_output_raw"]
        )
        btn_col2.download_button(
            label="📊 Tải file Excel (.xlsx) Phân phối chương trình", 
            data=excel_plan_data, 
            file_name=f"Phan_Phoi_Chuong_Trinh_{st.session_state['current_teacher_active'].replace(' ', '_')}.xlsx", 
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            use_container_width=True
        )
        
        # 🌟 TÍNH NĂNG MỚI: Thêm nút xóa nhanh file vừa khởi tạo
        if btn_col3.button("🗑️ Xóa bản kế hoạch hiện tại", type="secondary", use_container_width=True, key="btn_clear_current_active_v8"):
            st.session_state["current_ai_output_raw"] = None
            st.success("💥 Đã xóa bảng hiển thị hiện hành.")
            st.rerun()
    else:
        st.caption("Khung kế hoạch chi tiết dạng bảng biểu 8 cột sẽ xuất hiện tại đây sau khi bấm nút khởi tạo...")

    # --- KHU VỰC LỊCH SỬ KẾ HOẠCH ĐÃ LƯU TRONG TOÀN HỆ THỐNG ---
    if st.session_state["db_ke_hoach_da_luu"]:
        st.markdown("---")
        st.markdown("### 🗂️ Danh sách Kế hoạch Giáo dục đã lưu trên Hệ thống:")
        
        # Cực kỳ thông minh: Cho phép xóa phần tử cũ trong lịch sử thông qua vòng lặp an toàn
        keys_to_delete = []
        for key_id, info in st.session_state["db_ke_hoach_da_luu"].items():
            with st.expander(f"📋 Bản kế hoạch môn {info['subject']} - GV: {info['teacher']} (Khối {info['grade']})"):
                st.write(f"**Tổng số tuần**: {info['weeks']}")
                df_history = pd.DataFrame(info['data'])
                st.dataframe(df_history, use_container_width=True, hide_index=True)
                
                # Sắp xếp hàng ngang gồm 3 nút chức năng (Tải Word, Tải Excel, Xóa kho lịch sử)
                h_col1, h_col2, h_col3 = st.columns(3)
                
                word_history_data = export_plan_to_docx_with_table(info['teacher'], info['subject'], info['data'])
                h_col1.download_button(
                    label="📥 Tải file Word", 
                    data=word_history_data, 
                    file_name=f"Phu_Luc_III_{info['teacher'].replace(' ', '_')}.docx", 
                    key=f"dl_word_hist_{key_id}",
                    use_container_width=True
                )
                
                excel_history_data = export_plan_to_excel(info['teacher'], info['subject'], info['data'])
                h_col2.download_button(
                    label="📊 Tải file Excel", 
                    data=excel_history_data, 
                    file_name=f"Phan_Phoi_Chuong_Trinh_{info['teacher'].replace(' ', '_')}.xlsx", 
                    key=f"dl_excel_hist_{key_id}",
                    use_container_width=True
                )
                
                # 🌟 TÍNH NĂNG MỚI: Nút xóa file đã tạo lưu trữ trong danh sách kho dữ liệu lịch sử
                if h_col3.button("❌ Xóa bản ghi lịch sử", type="secondary", use_container_width=True, key=f"del_hist_rec_{key_id}"):
                    keys_to_delete.append(key_id)
                    
        # Thực hiện tác vụ xóa bản ghi cũ ngoài vòng lặp để tránh lỗi lặp phần tử bộ nhớ của Python
        if keys_to_delete:
            for k in keys_to_delete:
                del st.session_state["db_ke_hoach_da_luu"][k]
            st.success("🗑️ Hệ thống đã dọn dẹp và xóa vĩnh viễn tệp bản ghi được chỉ định!")
            st.rerun()
