import streamlit st
import sqlite3
import pandas as pd
import io

DB_PATH = "teacher_assistant.db"

# --- KHỞI TẠO BẢNG LƯU TRỮ BIÊN BẢN HỌP ---
def setup_minutes_database():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS org_minutes (
        id INTEGER PRIMARY KEY AUTOINCREMENT, 
        meeting_date TEXT, 
        session_number TEXT UNIQUE,
        present_members TEXT, 
        absent_members TEXT, 
        content_text TEXT, 
        resolution TEXT
    )
    """)
    conn.commit()
    conn.close()

# --- HÀM TRÍCH XUẤT VĂN BẢN TỪ FILE TẢI LÊN ---
def extract_text_from_minutes_upload(uploaded_file):
    import docx
    from pypdf import PdfReader
    text_content = ""
    try:
        if uploaded_file.name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            for p in doc.paragraphs:
                if p.text: text_content += p.text + "\n"
        elif uploaded_file.name.endswith('.pdf'):
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                text_content += (page.extract_text() or "") + "\n"
        elif uploaded_file.name.endswith('.txt'):
            text_content = uploaded_file.read().decode("utf-8")
    except Exception as e:
        st.error(f"Lỗi bóc tách file: {e}")
    return text_content.strip()
# --- HÀM ĐÓNG GÓI XUẤT FILE WORD BIÊN BẢN CHUẨN CĂN LỀ ---
def export_minutes_to_docx(meeting_date, session_number, present_members, absent_members, content_text, resolution):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
        
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n").bold = True
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(f"\nBIÊN BẢN SINH HOẠT TỔ CHUYÊN MÔN\n(Số: {session_number})\n")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(255, 0, 0)
    
    items = [
        f"- Ngày tiến hành họp tổ: {meeting_date}", 
        f"- Thành phần tham dự: {present_members}", 
        f"- Vắng mặt (Lý do): {absent_members}", 
        "\n--- DIỄN BIẾN CHI TIẾT CUỘC HỌP ---", 
        content_text, 
        "\n--- NGHỊ QUYẾT / KẾT LUẬN CHUNG ---", 
        resolution
    ]
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(item)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        if "---" in item: r.bold = True
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- GIAO DIỆN CHÍNH PHÂN HỆ BIÊN BẢN HỌP TỔ ---
def render_meeting_minutes():
    setup_minutes_database()
    st.markdown("<h3 style='text-align: left; color: #1E3A8A;'>📝 TRÌNH LẬP BIÊN BẢN HỌP TỔ CHUYÊN MÔN NÂNG CAO</h3>", unsafe_allow_html=True)
    
    tab_lap, tab_kho = st.tabs(["✍️ Lập biên bản mới", "🗄️ Kho lưu trữ biên bản họp"])
    
    if "minutes_content_draft" not in st.session_state: st.session_state["minutes_content_draft"] = ""
    if "minutes_resolution_draft" not in st.session_state: st.session_state["minutes_resolution_draft"] = ""

    # ==================== THÈ 1: LẬP BIÊN BẢN MỚI ====================
    with tab_lap:
        st.markdown("**📁 Nạp file tài liệu thảo luận / Kế hoạch tuần (Chèn tự động vào biên bản):**")
        file_ke_hoach = st.file_uploader("Kéo thả file kế hoạch (.docx, .pdf, .txt):", type=["docx", "pdf", "txt"], key="file_minutes_uploader_fixed")
        
        if file_ke_hoach and st.button("🧠 AI: Nghiên cứu file & Tóm tắt chèn vào diễn biến cuộc họp", type="secondary", use_container_width=True, key="btn_ai_extract_plan"):
            with st.spinner("AI đang trích xuất dữ liệu học liệu số..."):
                raw_text = extract_text_from_minutes_upload(file_ke_hoach)
                if raw_text:
                    try:
                        from app import run_ai_prompt_safe
                        api_key_system = st.secrets.get("GEMINI_API_KEY", "")
                        res_text, _ = run_ai_prompt_safe(f"Hãy tóm tắt văn bản này thành diễn biến một buổi sinh hoạt biên bản họp tổ chuyên môn cực kỳ chi tiết, chia các đầu mục rõ ràng, loại bỏ toàn bộ dấu sao sao **, chỉ dùng dấu gạch ngang '-': {raw_text}", api_key_system)
                        st.session_state["minutes_content_draft"] = res_text
                        st.success("✅ Đã nghiên cứu tài liệu và chèn dữ liệu tự động thành công!")
                        st.rerun()
                    except Exception as e: 
                        st.error(f"Lỗi kết nối AI: {e}")

        with st.form("form_create_minutes_fixed", border=False):
            c1, c2 = st.columns(2)
            m_date = c1.date_input("Ngày họp:")
            m_num = c2.text_input("Số biên bản (Ký hiệu):", placeholder="Ví dụ: BB-02/KHTN")
            
            m_present = st.text_area("Thành phần tham dự (Ghi rõ họ tên):", placeholder="Thầy Dưỡng (Chủ trì), Cô Trang (Thư ký)...")
            m_absent = st.text_input("Vắng mặt (Lý do):", placeholder="Không")
            
            st.markdown("**📄 Diễn biến cuộc họp và Ý kiến đóng góp (Đã liên kết chèn tự động):**")
            m_content = st.text_area("Nội dung chi tiết diễn biến:", value=st.session_state["minutes_content_draft"], height=200, key="txt_area_m_content")
            
            st.markdown("🤖 **Hỗ trợ đúc kết Kết luận từ Trợ lý AI:**")
            col_ai_btn, _ = st.columns([2.0, 2.0])
            with col_ai_btn:
                run_ai_res = st.form_submit_button("🧠 AI: Tóm tắt Nghị quyết cuộc họp dựa trên diễn biến ở trên")
                
            if run_ai_res:
                if m_content:
                    with st.spinner("AI đang đúc kết nghị quyết quyết nghị..."):
                        try:
                            from app import run_ai_prompt_safe
                            api_key_system = st.secrets.get("GEMINI_API_KEY", "")
                            res_ai, _ = run_ai_prompt_safe(f"Hãy dựa trên văn bản diễn biến cuộc họp tổ sau đây để viết ra các mục nghị quyết/kết luận ngắn gọn chung, không dùng dấu sao sao **, chỉ dùng dấu gạch ngang đầu dòng liệt kê: {m_content}", api_key_system)
                            st.session_state["minutes_resolution_draft"] = res_ai
                            st.rerun()
                        except: pass
                else:
                    st.warning("⚠️ Vui lòng nhập nội dung diễn biến cuộc họp trước!")

            m_resolution = st.text_area("Nghị quyết / Kết luận chung cuộc họp (Thầy có thể chỉnh sửa):", value=st.session_state["minutes_resolution_draft"], height=120, key="txt_area_m_resolution")
            st.markdown("<br>", unsafe_allow_html=True)
            save_minutes = st.form_submit_button("💾 KHÓA & LƯU BIÊN BẢN VÀO HỆ THỐNG", type="primary", use_container_width=True)
        if save_minutes:
            if not m_num or not m_content:
                st.warning("⚠️ Vui lòng điền đầy đủ thông tin Số biên bản và Diễn biến cuộc họp!")
            else:
                conn = sqlite3.connect(DB_PATH)
                conn.execute("INSERT OR REPLACE INTO org_minutes (meeting_date, session_number, present_members, absent_members, content_text, resolution) VALUES (?, ?, ?, ?, ?, ?)", (str(m_date), m_num.strip(), m_present.strip(), m_absent.strip(), m_content.strip(), m_resolution.strip()))
                conn.commit()
                conn.close()
                st.success(f"🎉 Biên bản số **{m_num}** đã được lưu trữ vĩnh viễn thành công!")
                st.session_state["minutes_content_draft"] = ""
                st.session_state["minutes_resolution_draft"] = ""
                st.rerun()

    # ==================== THÈ 2: KHO LƯU TRỮ BIÊN BẢN HỌP TỔ ====================
    with tab_kho:
        st.markdown("#### 🗄️ Thư viện lưu trữ biên bản sinh hoạt tổ")
        conn = sqlite3.connect(DB_PATH)
        all_minutes = conn.execute("SELECT id, meeting_date, session_number, present_members, absent_members, content_text, resolution FROM org_minutes ORDER BY meeting_date DESC").fetchall()
        conn.close()
        
        if not all_minutes:
            st.info("Hiện tại hệ thống chưa lưu trữ biên bản cuộc họp nào.")
        else:
            for idx, row in enumerate(all_minutes):
                m_id, m_dt, m_snum, m_pres, m_abs, m_ct, m_resol = row
                col_exp, col_del = st.columns([0.88, 0.12])
                with col_exp:
                    with st.expander(f"📝 {idx+1}. Biên bản số: {m_snum} - Ngày họp: {m_dt}"):
                        st.write(f"**👥 Thành phần:** {m_pres} | **❌ Vắng:** {m_abs}")
                        st.caption(f"**📄 Diễn biến cuộc họp:**\n{m_ct}")
                        st.write(f"**🏅 Nghị quyết kết luận:**\n{m_resol}")
                        
                        # Xuất file Word chuẩn cho từng cuộc họp độc lập tĩnh khóa Key
                        w_data = export_minutes_to_docx(m_dt, m_snum, m_pres, m_abs, m_ct, m_resol)
                        st.download_button(label="📥 Tải file Word (.docx) biên bản này", data=w_data, file_name=f"Bien_Ban_So_{m_snum.replace('/', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_w_minutes_{m_id}", use_container_width=True)
                with col_del:
                    st.write("")
                    if st.button("🗑️ Xóa", key=f"del_m_minutes_{m_id}", use_container_width=True):
                        conn = sqlite3.connect(DB_PATH)
                        conn.execute("DELETE FROM org_minutes WHERE id=?", (m_id,))
                        conn.commit()
                        conn.close()
                        st.success("Đã xóa!")
                        st.rerun()
