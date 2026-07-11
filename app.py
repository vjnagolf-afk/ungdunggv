import streamlit as st
import io
import json
import os
import pandas as pd
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from pypdf import PdfReader

# ==============================================================================
# 1. CẤU HÌNH TRANG VÀ THIẾT LẬP CSS GIAO DIỆN SƯ PHẠM ĐẸP MẮT
# ==============================================================================
st.set_page_config(
    page_title="EduAssist AI 2.0 - Hệ Sinh Thái Giáo Viên Số",
    page_icon="🚀",
    layout="wide"
)

st.markdown("""
<style>
/* Đẩy nội dung thanh bên lên sát đỉnh đầu một cách vừa vặn */
div[data-testid="stSidebarUserContent"] {
    padding-top: 1.0rem !important;
    padding-bottom: 0px !important;
}
/* Rút bớt lề của đường kẻ phân cách để hai khối nội dung sát lại nhau hơn */
div[data-testid="stSidebarUserContent"] hr {
    margin-top: 0.6rem !important;
    margin-bottom: 0.6rem !important;
}
/* Thiết lập khoảng cách dọc giữa các Widget ở mức vừa phải, đẹp mắt */
div[data-testid="stSidebarUserContent"] div[data-testid="stVerticalBlock"] > div {
    margin-bottom: 0px !important;
}
/* Hiệu ứng rê chuột mượt mà cho tất cả các nút bấm */
div.stButton > button, div.stDownloadButton > button {
    transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1) !important;
    border-radius: 8px !important;
}
div.stButton > button:hover, div.stDownloadButton > button:hover {
    background-color: #1E40AF !important; 
    color: #FFFFFF !important; 
    border-color: #3B82F6 !important; 
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 15px rgba(59, 130, 246, 0.4) !important;
}
</style>
""", unsafe_allow_html=True)
# ==============================================================================
# 2. BỘ ENGINE AI TOÀN DIỆN CHO HỆ THỐNG
# ==============================================================================
def run_ai_prompt_safe(prompt_text, api_key):
    models_to_try = ["gemini-2.5-flash", "gemini-1.5-flash", "gemini-2.5-pro"]
    client = genai.Client(api_key=api_key)
    last_error = None
    for model in models_to_try:
        for attempt in range(2):
            try:
                response = client.models.generate_content(model=model, contents=prompt_text)
                return response.text, model
            except Exception as e:
                last_error = e
                time.sleep(2)
    raise Exception(f"Tất cả các mô hình AI đều đang gặp lỗi: {last_error}")
# ==============================================================================
# 3. BỘ ENGINE KẾT XUẤT VĂN BẢN CHUẨN ĐỊNH DẠNG HÀNH CHÍNH
# ==============================================================================
def export_to_docx_vietnam_standard(text_content, title_name, school_name="TRƯỜNG THCS NGUYỄN CHÍ THANH", group_name="TỔ KHOA HỌC TỰ NHIÊN - GDTC"):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    admin_table = doc.add_table(rows=1, cols=2)
    admin_table.autofit = False
    admin_table.columns.width = Inches(3.2)
    admin_table.columns.width = Inches(3.8)
    
    cell_l = admin_table.rows.cells
    p_left = cell_l.paragraphs
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.add_run(f"{school_name.upper()}\n").bold = True
    p_left.add_run(f"{group_name.upper()}\n").bold = True
    p_left.add_run("Số: ..... /BB-TCM").font.size = Pt(11)
    
    cell_r = admin_table.rows.cells
    p_right = cell_r.paragraphs
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").bold = True
    p_right.add_run("Độc lập - Tự do - Hạnh phúc\n").bold = True
    p_right.add_run("***************").font.size = Pt(11)
    
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run(title_name.upper()).bold = True
    
    for line in text_content.split('\n'):
        cleaned_line = line.strip()
        if not cleaned_line: continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if cleaned_line.startswith('#'):
            p.add_run(cleaned_line.lstrip('#').strip()).bold = True
        else:
            p.add_run(cleaned_line)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def convert_df_to_excel_bytes(df, sheet_name='Sheet1'):
    val_output = io.BytesIO()
    with pd.ExcelWriter(val_output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return val_output.getvalue()
def generate_dynamic_5512_docx():
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("KHUNG KẾ HOẠCH BÀI DẠY MẪU CHUẨN 5512").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def read_uploaded_docx(uploaded_file):
    try:
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    except: return "Lỗi đọc file Word"

def read_uploaded_pdf(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except: return "Lỗi đọc file PDF"

# Khởi tạo cơ sở dữ liệu gốc ban đầu của tổ chuyên môn
if "db_thanh_vien" not in st.session_state:
    st.session_state["db_thanh_vien"] = [
        {"STT": 1, "Họ và tên": "Lê Hồng Dưỡng", "Chức vụ": "Tổ trưởng", "Phân môn chính": "Vật lý", "Email": "duonglh@gmail.com", "Số điện thoại": "0912345678"},
        {"STT": 2, "Họ và tên": "Nguyễn Thị Bình", "Chức vụ": "Tổ phó", "Phân môn chính": "Sinh học", "Email": "binhnt@gmail.com", "Số điện thoại": "0987654321"},
        {"STT": 3, "Họ và tên": "Trần Thanh Việt", "Chức vụ": "Giáo viên", "Phân môn chính": "Hóa học", "Email": "viettt@gmail.com", "Số điện thoại": "0905123456"}
    ]

DANH_SACH_NAM_HOC = [f"{nam} - {nam+1}" for nam in range(2020, 2036)]

if "db_de_kiem_tra" not in st.session_state:
    st.session_state["db_de_kiem_tra"] = [
        {
            "ten_de": "Đề kiểm tra giữa kỳ I - KHTN 9 (Mẫu tham khảo)",
            "mon": "Khoa học tự nhiên",
            "khoi": "Lớp 9",
            "noi_dung": "### MA TRẬN & ĐẶC TẢ ĐỀ KIỂM TRA\\n| Chủ đề | Nhận biết (TNKQ) | Thông hiểu (TNKQ) | Vận dụng (TL) | Vận dụng cao (TL) | Tổng số câu | Điểm |\\n| :--- | :---: | :---: | :---: | :---: | :---: | :---: |\\n| 1. Thấu kính | 4 câu | 2 câu | 1 câu | 0 | 7 câu | 3.5 |\\n| **Tổng cộng** | **4 câu** | **2 câu** | **1 câu** | **0** | **7 câu** | **10.0** |"
        }
    ]
if "db_thanh_tich_da_nam" not in st.session_state:
    st.session_state["db_thanh_tich_da_nam"] = {}
    st.session_state["db_thanh_tich_da_nam"]["2025 - 2026"] = [
        {"Năm học": "2025 - 2026", "Họ và tên": "Lê Hồng Dưỡng", "Đánh giá viên chức": "Hoàn thành xuất sắc (HTSX)", "Kết quả BD HSG": "1 giải Nhất môn Vật lý", "Kết quả NCKH": "Đạt giải Nhì", "Kết quả STEM": "02 bài học xuất sắc", "Kết quả Sáng kiến": "Giải B", "Thi GVDG": "Cấp Tỉnh", "Thi GVCNG": "Không", "Kết quả TDTT": "Giải Nhất bóng bàn", "Kết quả VN": "Giải xuất sắc", "Hiến máu nhân đạo": "01 lần", "Hoạt động khác": "Tích cực Chuyển đổi số"},
        {"Năm học": "2025 - 2026", "Họ và tên": "Nguyễn Thị Bình", "Đánh giá viên chức": "Hoàn thành tốt (HTT)", "Kết quả BD HSG": "2 giải Ba môn Sinh", "Kết quả NCKH": "Cấp Trường", "Kết quả STEM": "01 chuyên đề Tốt", "Kết quả Sáng kiến": "Cấp Huyện", "Thi GVDG": "Cấp Huyện", "Thi GVCNG": "Không", "Kết quả TDTT": "Đội trưởng Bóng chuyền", "Kết quả VN": "Biên đạo giải Nhất", "Hiến máu nhân đạo": "01 lần", "Hoạt động khác": "Công đoàn viên xuất sắc"}
    ]

def sync_sub_databases():
    current_pc = st.session_state.get("db_phan_cong_hien_tai", [])
    pc_mapped = {x["Họ và tên"]: x for x in current_pc}
    new_pc = []
    for tv in st.session_state["db_thanh_vien"]:
        name = tv["Họ và tên"]
        if name in pc_mapped:
            old_item = pc_mapped[name]
            new_pc.append({
                "STT": int(tv["STT"]), "Họ và tên": name, "Phân môn chính": tv["Phân môn chính"],
                "Lớp phân công": old_item.get("Lớp phân công", "9A1, 9A2"), "Số tiết phân công": old_item.get("Số tiết phân công", 8),
                "Nhiệm vụ kiêm nhiệm": tv["Chức vụ"]
            })
        else:
            new_pc.append({
                "STT": int(tv["STT"]), "Họ và tên": name, "Phân môn chính": tv["Phân môn chính"],
                "Lớp phân công": "", "Số tiết phân công": 0, "Nhiệm vụ kiêm nhiệm": tv["Chức vụ"]
            })
    st.session_state["db_phan_cong_hien_tai"] = new_pc

    for nam_hoc in DANH_SACH_NAM_HOC:
        if nam_hoc not in st.session_state["db_thanh_tich_da_nam"]:
            st.session_state["db_thanh_tich_da_nam"][nam_hoc] = []
        current_tt = st.session_state["db_thanh_tich_da_nam"][nam_hoc]
        tt_mapped = {x["Họ và tên"]: x for x in current_tt}
        new_tt = []
        for tv in st.session_state["db_thanh_vien"]:
            name = tv["Họ và tên"]
            if name in tt_mapped:
                new_tt.append(tt_mapped[name])
            else:
                new_tt.append({
                    "Năm học": nam_hoc, "Họ và tên": name, "Đánh giá viên chức": "Hoàn thành tốt (HTT)",
                    "Kết quả BD HSG": "Không", "Kết quả NCKH": "Không", "Kết quả STEM": "Không", "Kết quả Sáng kiến": "Không",
                    "Thi GVDG": "Không", "Thi GVCNG": "Không", "Kết quả TDTT": "Không", "Kết quả VN": "Không", "Hiến máu nhân đạo": "Không", "Hoạt động khác": "Không"
                })
        st.session_state["db_thanh_tich_da_nam"][nam_hoc] = new_tt

if "db_phan_cong_hien_tai" not in st.session_state:
    sync_sub_databases()

has_secrets = "GEMINI_API_KEY" in st.secrets if hasattr(st, "secrets") else False
api_key_input = st.secrets["GEMINI_API_KEY"] if has_secrets else ""

st.sidebar.title("🛠️ Menu Tiện Ích 2.0")
phan_he_lam_viec = st.sidebar.radio("CHỌN PHÂN HỆ TÁC NGHIỆP", [" Trợ lý Giảng dạy (Giáo viên)", " Trợ lý Quản lý (Tổ chuyên môn)"], index=0)

if phan_he_lam_viec == " Trợ lý Giảng dạy (Giáo viên)":
    chuc_nang_chinh = st.sidebar.selectbox("CHỌN NỘI DUNG THỰC HIỆN", ["1. Thiết kế KHBD thông minh", "2. Thiết kế Đề KTĐG (Ma trận - Đặc tả)", "3. Đánh giá học sinh"])
else:
    chuc_nang_chinh = st.sidebar.selectbox(
        "QUẢN LÝ TỔ CHUYÊN MÔN", 
        [
            "1. Hệ thống Quản lý và Phân công chuyên môn giảng dạy", 
            "2. Xây dựng Biên bản sinh hoạt tổ chuyên môn định kỳ",
            "3. Xây dựng Kế hoạch Giáo dục cá nhân (Phụ lục III - Công văn 5512)",
            "4. Thống kê số liệu tổ"
        ]
    )

st.sidebar.markdown("---")
st.sidebar.title("🔐 Xác thực Quyền Truy cập")
user_role = st.sidebar.selectbox("Bạn đăng nhập với tư cách:", ["Giáo viên tổ viên", "Tổ trưởng chuyên môn (Admin)"], index=1)
is_admin = (user_role == "Tổ trưởng chuyên môn (Admin)")

if not is_admin: st.sidebar.info("💡 Chế độ Tổ viên: Bạn được xem toàn bộ hồ sơ tổ.")
else: st.sidebar.success("🔑 Quyền Admin: Thầy Dưỡng có toàn quyền cập nhật.")

if not has_secrets:
    api_key_input = st.sidebar.text_input("Nhập khóa Gemini API Key nếu cần dùng AI:", type="password", value=api_key_input)

st.sidebar.markdown("---")
tac_gia = st.sidebar.text_input("Họ và tên tác giả:", value="Lê Hồng Dưỡng")
don_vi = st.sidebar.text_input("Đơn vị công tác:", value="Trường THCS Nguyễn Chí Thanh")
# Giao diện chính phân hệ Giáo viên
if phan_he_lam_viec == " Trợ lý Giảng dạy (Giáo viên)":
    if chuc_nang_chinh == "1. Thiết kế KHBD thông minh":
        st.header("📋 Thiết kế Kế hoạch bài dạy (KHBD) tích hợp năng lực số")
        st.markdown("##### 📂 Kho tài liệu hướng dẫn xây dựng KHBD")
        st.info("Thầy cô có thể tải trực tiếp tài liệu hướng dẫn cấu trúc khung kế hoạch bài dạy mẫu 5512 thế hệ mới:")
        
        col_w, col_t = st.columns(2)
        with col_w: st.download_button(label="📄 Tải Khung KHBD chuẩn (.docx)", data=generate_dynamic_5512_docx(), file_name="Khung_Ke_hoach_Bai_day_5512.docx")
        with col_t: st.download_button(label="📥 Tải Cẩm nang sư phạm (.txt)", data="CẨM NANG SƯ PHẠM 5512 NÂNG CAO".encode('utf-8'), file_name="Cam_nang_Huong_dan_5512.txt")
        
        st.markdown("---")
        col_input1, col_input2 = st.columns(2)
        with col_input1:
            ten_bai = st.text_input("Tên bài học học tập:", placeholder="Ví dụ: Bài 1: Qua Đèo Ngang")
            lop = st.selectbox("Khối lớp giảng dạy:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"], index=3)
        with col_input2:
            mon_hoc = st.text_input("Môn học chuyên môn:", value="Khoa học tự nhiên")
            thoi_luong = st.text_input("Thời lượng thực hiện:", value="1 tiết")
            
        uploaded_khbd_file = st.file_uploader("Tải file học liệu gốc/Nội dung SGK nguồn (.pdf, .docx):", type=["pdf", "docx"], key="khbd_up")
        khbd_content = ""
        if uploaded_khbd_file:
            khbd_content = read_uploaded_docx(uploaded_khbd_file) if uploaded_khbd_file.name.endswith('.docx') else read_uploaded_pdf(uploaded_khbd_file)
            st.success("✅ Đã nạp học liệu gốc thành công vào bộ nhớ AI.")
            
        uu_tien_tai_lieu = st.checkbox("🎯 ƯU TIÊN BÁM SÁT 100% TÀI LIỆU TẢI LÊN", value=True)
        if st.button("🚀 XD KHBD thông minh"):
            with st.spinner("AI đang dựng giáo án..."):
                try:
                    res, _ = run_ai_prompt_safe(f"Soạn giáo án chuẩn 5512 cho bài {ten_bai}", api_key_input)
                    st.markdown(res)
                except Exception as e: st.error(f"Lỗi: {e}")
    elif chuc_nang_chinh == "2. Thiết kế Đề KTĐG (Ma trận - Đặc tả)":
        st.header("📊 Thiết kế Hệ thống Đề kiểm tra đánh giá định kỳ")
        tab_thiet_ke, tab_kho_luu_tru = st.tabs(["✨ Thiết kế đề thi mới", "📂 Thư mục lưu trữ đề đã dựng"])
        
        with tab_thiet_ke:
            col_de1, col_de2 = st.columns(2)
            with col_de1:
                ten_de = st.text_input("Tên kỳ kiểm tra thiết lập:", placeholder="Ví dụ: Kiểm tra định kỳ học kỳ I")
                mon_de = st.text_input("Môn học kiểm tra:", value="Khoa học tự nhiên")
                khoi_de = st.selectbox("Khối lớp đề thi:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"], index=3)
            with col_de2:
                thoi_gian_de = st.text_input("Thời gian làm bài thi:", value="45 phút")
                ty_le_de = st.text_input("Tỷ lệ ma trận mong muốn:", value="Nhận biết: 40% - Thông hiểu: 30% - Vận dụng thấp: 20% - Vận dụng cao: 10%")
                
            st.markdown("#### ⚙️ Cấu hình cấu trúc câu hỏi chi tiết")
            col_sl1, col_sl2 = st.columns(2)
            with col_sl1:
                tong_so_tn = st.slider("Tổng số câu trắc nghiệm khách quan (Tối đa 32 câu):", min_value=0, max_value=32, value=16)
            with col_sl2:
                tong_so_tl = st.number_input("Tổng số câu hỏi tự luận tùy chọn (Không giới hạn):", min_value=0, value=2, step=1)
                
            # Khôi phục nguyên vẹn khung phân bổ chi tiết hình thức trắc nghiệm
            st.markdown("##### 🔍 Phân bổ chi tiết hình thức câu hỏi trắc nghiệm")
            col_hb1, col_sl_cs2, col_sl_cs3, col_sl_cs4 = st.columns(4)
            with col_hb1:
                tn_1_dap_an = st.number_input("Trắc nghiệm 1 đáp án đúng:", min_value=0, value=10, step=1)
            with col_sl_cs2:
                tn_dung_sai = st.number_input("Trắc nghiệm Đúng/Sai:", min_value=0, value=2, step=1)
            with col_sl_cs3:
                tn_dien_khuyen = st.number_input("Trắc nghiệm điền khuyết:", min_value=0, value=2, step=1)
            with col_sl_cs4:
                tn_tra_loi_ngan = st.number_input("Trắc nghiệm trả lời ngắn:", min_value=0, value=2, step=1)
                
            # Trạng thái đồng bộ tự động kiểm tra số lượng câu trắc nghiệm khớp ảnh chụp
            tong_phan_bo_thuc_te = tn_1_dap_an + tn_dung_sai + tn_dien_khuyen + tn_tra_loi_ngan
            if tong_phan_bo_thuc_te == tong_so_tn:
                st.success(f"✅ Đồng bộ thành công! Tổng phân bổ ({tong_phan_bo_thuc_te} câu) khớp với cấu hình đề thi.")
            else:
                st.warning(f"⚠️ Chưa đồng bộ! Tổng phân bổ chi tiết ({tong_phan_bo_thuc_te} câu) chưa khớp với Tổng số câu trắc nghiệm đặt trên thanh trượt ({tong_so_tn} câu).")
                
            uploaded_excel_de = st.file_uploader("Tải lên file giới hạn kiến thức hoặc bộ đề thi làm căn cứ (.pdf, .docx):", type=["pdf", "docx"], key="up_file_de_thi")
            content_de_nguon = ""
            if uploaded_excel_de:
                content_de_nguon = read_uploaded_docx(uploaded_excel_de) if uploaded_excel_de.name.endswith('.docx') else read_uploaded_pdf(uploaded_excel_de)
                st.success("✅ Đã nạp thành công tài liệu kiến thức nguồn vào bộ nhớ AI.")
                
            uu_tien_de = st.checkbox("🔥 ƯU TIÊN BÁM SÁT 100% TÀI LIỆU/ĐỀ MẪU TẢI LÊN", value=True)
            if st.button("⚡ Thiết lập Đề thi + Ma trận & Đặc tả"):
                if not api_key_input: 
                    st.error("Thầy cần cấu hình Gemini API Key tại thanh bên!")
                else:
                    with st.spinner("AI đang dựng bảng ma trận, bảng đặc tả và hệ thống câu hỏi chi tiết..."):
                        try:
                            constraint_text = f"BẮT BUỘC bám sát nội dung tài liệu này:\n{content_de_nguon}\n" if (uu_tien_de and content_de_nguon) else ""
                            prompt_de = f"""
                            Hãy đóng vai là một chuyên gia khảo thí Việt Nam xuất sắc.
                            {constraint_text}
                            Thiết kế một đề kiểm tra hoàn chỉnh môn {mon_de} thuộc khối {khoi_de}. Tên kỳ kỳ: {ten_de}. Tỷ lệ: {ty_le_de}.
                            CẤU TRÚC ĐỀ BẮT BUỘC:
                            - Phần Trắc nghiệm: Tổng {tong_so_tn} câu. Trong đó gồm: {tn_1_dap_an} câu 4 lựa chọn, {tn_dung_sai} câu Đúng/Sai, {tn_dien_khuyen} câu điền khuyết, {tn_tra_loi_ngan} câu trả lời ngắn.
                            - Phần Tự luận: Gồm {tong_so_tl} câu hỏi tự luận.
                            ĐẦU RA GỒM: 1. BẢNG MA TRẬN ĐỀ, 2. BẢNG ĐẶC TẢ ĐỀ, 3. ĐỀ KIỂM TRA, 4. ĐÁP ÁN CHI TIẾT.
                            """
                            result_text, _ = run_ai_prompt_safe(prompt_de, api_key_input)
                            st.markdown(result_text)
                            st.session_state["db_de_kiem_tra"].append({
                                "ten_de": ten_de if ten_de else "Đề thi tự động tạo mới",
                                "mon": mon_de, "khoi": khoi_de, "noi_dung": result_text
                            })
                            st.success("✅ Đã lưu bộ đề thi thành công vào Thư mục nội bộ!")
                            st.download_button(label="📥 Tải bộ đề kiểm tra bản Word (.docx)", data=export_to_docx_vietnam_standard(result_text, ten_de if ten_de else "De_Kiem_Tra"), file_name=f"{ten_de}.docx")
                        except Exception as e: st.error(f"Lỗi: {e}")
                            
        with tab_kho_luu_tru:
            st.subheader("📁 Thư mục lưu trữ đề kiểm tra nội bộ đã dựng")
            for idx, item in enumerate(st.session_state["db_de_kiem_tra"]):
                with st.expander(f"📋 {item['ten_de']} - Môn: {item['mon']} ({item['khoi']})"):
                    st.markdown(item["noi_dung"])
                    st.download_button(label="📥 Tải lại File Word", data=export_to_docx_vietnam_standard(item["noi_dung"], item["ten_de"]), file_name=f"{item['ten_de']}_{idx}.docx", key=f"dl_de_thi_{idx}")

    elif chuc_nang_chinh == "3. Đánh giá học sinh":
        st.header("💯 Bộ công cụ đánh giá, nhận xét và bổ trợ học sinh")
        chu_de = st.text_input("Nội dung đánh giá:", value="Quang học lớp 9")
        if st.button(" Sinh học liệu"):
            with st.spinner("AI đang tạo..."):
                try:
                    res, _ = run_ai_prompt_safe(f"Tạo tiêu chí rubric cho chủ đề {chu_de}", api_key_input)
                    st.markdown(res)
                except Exception as e: st.error(f"Lỗi: {e}")
else:
    # PHÂN HỆ: QUẢN LÝ TỔ CHUYÊN MÔN
    if chuc_nang_chinh == "1. Hệ thống Quản lý và Phân công chuyên môn giảng dạy":
        st.subheader("🎖️ Hệ thống Quản lý và Phân công chuyên môn giảng dạy")
        tab_thanh_vien, tab_phan_cong, tab_thanh_tich = st.tabs(["👥 Danh sách thành viên & Chuyên môn", "📅 Lập & Quản lý Phân công chuyên môn (Đợt)", "📈 Bảng theo dõi Thành tích Giáo viên"])
        
        with tab_thanh_vien:
            df_thanh_vien = pd.DataFrame(st.session_state["db_thanh_vien"])
            df_tv_edited = st.data_editor(df_thanh_vien, num_rows="dynamic" if is_admin else "fixed", disabled=not is_admin, use_container_width=True, key="editor_thanh_vien")
            if is_admin:
                col_save_tv, col_ex_tv = st.columns(2)
                with col_save_tv:
                    if st.button("💾 Lưu cập nhật Danh sách & Tự động Đồng bộ các Tab"):
                        st.session_state["db_thanh_vien"] = df_tv_edited.to_dict('records')
                        sync_sub_databases()
                        st.success("✅ Đã cập nhật dữ liệu gốc và đồng bộ nhân sự toàn hệ thống thành công!")
                        st.rerun()
                with col_ex_tv:
                    df_mau_tv = pd.DataFrame([{"STT": 1, "Họ và tên": "Nguyễn Văn A", "Chức vụ": "Giáo viên", "Phân môn chính": "Vật lý", "Email": "anv@gmail.com", "Số điện thoại": "0900000000"}])
                    st.download_button(label="📥 Tải File Excel Cấu trúc Mẫu (.xlsx)", data=convert_df_to_excel_bytes(df_mau_tv), file_name="Mau_Danh_Sach_Thanh_Vien.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                st.markdown("---")
                uploaded_excel_tv = st.file_uploader("📥 Nạp danh sách thành viên nhanh từ file Excel (.xlsx):", type=["xlsx"], key="up_excel_tv")
                if uploaded_excel_tv:
                    try:
                        df_uploaded_tv = pd.read_excel(uploaded_excel_tv)
                        required_cols = ["STT", "Họ và tên", "Chức vụ", "Phân môn chính", "Email", "Số điện thoại"]
                        for col in required_cols:
                            if col not in df_uploaded_tv.columns: df_uploaded_tv[col] = ""
                        df_uploaded_tv["STT"] = pd.to_numeric(df_uploaded_tv["STT"], errors='coerce').fillna(0).astype(int)
                        st.session_state["db_thanh_vien"] = df_uploaded_tv[required_cols].to_dict('records')
                        sync_sub_databases()
                        st.success("🎉 Đã nhập danh sách từ Excel!")
                        st.rerun()
                    except Exception as e: st.error(f"Lỗi tệp Excel: {e}")

        with tab_phan_cong:
            df_pc_current = pd.DataFrame(st.session_state["db_phan_cong_hien_tai"])
            df_pc_edited = st.data_editor(df_pc_current, num_rows="fixed", disabled=not is_admin, use_container_width=True, key="editor_phan_cong_active")
            if is_admin and st.button("💾 Lưu dữ liệu Phân công giảng dạy"):
                st.session_state["db_phan_cong_hien_tai"] = df_pc_edited.to_dict('records')
                st.success("✅ Đã lưu phân công!")
                st.rerun()

        with tab_thanh_tich:
            nam_hoc_selected = st.selectbox("📅 CHỌN NĂM HỌC CẦN QUẢN LÝ VÀ LƯU TRỮ HỒ SƠ THI ĐUA:", DANH_SACH_NAM_HOC, index=5)
            df_tt_current = pd.DataFrame(st.session_state["db_thanh_tich_da_nam"][nam_hoc_selected])
            df_tt_edited = st.data_editor(df_tt_current, num_rows="fixed", disabled=not is_admin, use_container_width=True, key=f"editor_thanh_tich_{nam_hoc_selected}")
            if is_admin:
                col_save_tt, col_ex_tt = st.columns(2)
                with col_save_tt:
                    if st.button(f"💾 Lưu dữ liệu Thi đua riêng cho năm {nam_hoc_selected}"):
                        st.session_state["db_thanh_tich_da_nam"][nam_hoc_selected] = df_tt_edited.to_dict('records')
                        st.success(f"✅ Hệ thống đã lưu trữ dữ liệu năm {nam_hoc_selected}!")
                        st.rerun()
                with col_ex_tt:
                    st.download_button(label=f"📥 Tải Bảng thi đua {nam_hoc_selected} về máy (.xlsx)", data=convert_df_to_excel_bytes(df_tt_current), file_name=f"Thi_Dua_Nam_Hoc_{nam_hoc_selected.replace(' - ', '_')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                st.markdown("---")
                uploaded_excel_tt = st.file_uploader(f"Tải tệp Excel kết quả thi đua của năm học {nam_hoc_selected} (.xlsx):", type=["xlsx"], key=f"up_excel_tt_{nam_hoc_selected}")
                if uploaded_excel_tt:
                    try:
                        df_uploaded_tt = pd.read_excel(uploaded_excel_tt)
                        required_tt_cols = ["Năm học", "Họ và tên", "Đánh giá viên chức", "Kết quả BD HSG", "Kết quả NCKH", "Kết quả STEM", "Kết quả Sáng kiến", "Thi GVDG", "Thi GVCNG", "Kết quả TDTT", "Kết quả VN", "Hiến máu nhân đạo", "Hoạt động khác"]
                        for col in required_tt_cols:
                            if col not in df_uploaded_tt.columns: df_uploaded_tt[col] = "Không"
                        df_uploaded_tt["Năm học"] = nam_hoc_selected
                        st.session_state["db_thanh_tich_da_nam"][nam_hoc_selected] = df_uploaded_tt[required_tt_cols].to_dict('records')
                        st.success(f"🎉 Đã nạp thành công tệp Excel vào niên khóa {nam_hoc_selected}!")
                        st.rerun()
                    except Exception as e: st.error(f"Lỗi file Excel thi đua: {e}")
    elif chuc_nang_chinh == "2. Xây dựng Biên bản sinh hoạt tổ chuyên môn định kỳ":
        st.subheader("📝 Xây dựng Biên bản sinh hoạt tổ chuyên môn (Chuẩn mẫu Trường)")
        st.markdown("#### I. THỜI GIAN, ĐỊA ĐIỂM, THÀNH PHẦN HOẠT ĐỘNG")
        col_t1, col_t2, col_t3 = st.columns(3)
        with col_t1:
            ky_hop_so = st.text_input("Kỳ họp thứ:", value="05")
            thang_hop_so = st.text_input("Biên bản họp tháng:", value="01")
            nam_hoc_chuoi = st.text_input("Năm học trường đăng ký:", value="2025 - 2026")
        with col_t2:
            gio_phut_hop = st.text_input("Vào lúc mấy giờ mấy phút:", value="13 giờ 30 phút")
            ngay_thang_hop = st.text_input("Ngày/Tháng/Năm cuộc họp diễn ra:", value="17/01/2026")
            phong_hop = st.text_input("Địa điểm tại phòng học:", value="Hội trường")
        with col_t3:
            chu_tri_cuoc = st.text_input("Chủ trì cuộc họp (Chủ tọa):", value="Thầy Lê Hồng Dưỡng – Tổ trưởng")
            thu_ky_cuoc = st.text_input("Thư ký lập biên bản:", value="Cô Nguyễn Thị Bình")
            thanh_phan_co = st.text_input("Thành phần (Có mặt / Vắng mặt):", value="Có mặt: 10/10; Vắng mặt: 0")

        st.markdown("---")
        st.markdown("#### II. NỘI DUNG CUỘC HỌP & NẠP KẾ HOẠCH NGUỒN")
        uploaded_kh_thang = st.file_uploader("📥 Tải file Kế hoạch tháng làm căn cứ biên bản (.pdf, .docx):", type=["pdf", "docx"], key="up_kh_thang")
        content_kh_thang = ""
        if uploaded_kh_thang:
            content_kh_thang = read_uploaded_docx(uploaded_kh_thang) if uploaded_kh_thang.name.endswith('.docx') else read_uploaded_pdf(uploaded_kh_thang)
            st.success("✅ Đã nạp thành công kế hoạch tháng vào bộ nhớ trợ lý ảo.")

        col_m1, col_m2, col_m3 = st.columns(3)
        with col_m1:
            bo_sung_11 = st.text_area("Ghi chú bổ sung mục 1.1 (Công tác trọng tâm):", placeholder="Nhập thêm lưu ý nếu có...")
        with col_m2:
            bo_sung_12 = st.text_area("Ghi chú bổ sung mục 1.2 (Thảo luận chuyên môn):", placeholder="Nhập lưu ý phương pháp, chuyên đề STEM...")
        with col_m3:
            bo_sung_13 = st.text_area("Ghi chú bổ sung mục 1.3 (Công tác khác):", placeholder="Nhập việc phong trào, công đoàn...")

        y_kien_dong_gop = st.text_area("Mục 2. Ý kiến đóng góp của các thành viên tổ chuyên môn:", value="- Các thành viên thảo luận sôi nổi và hoàn toàn nhất trí với phương hướng phân bổ chuyên môn tháng này.")

        if st.button("🤖 AI Tự động soạn thảo Biên bản"):
            if not uploaded_kh_thang: st.warning("⚠️ Thầy cần đính kèm file Kế hoạch tháng lên trước!")
            else:
                with st.spinner("AI đang dựng biên bản hành chính..."):
                    try:
                        prompt_bb = f"Hãy đóng vai thư ký tổ chuyên môn trường THCS Nguyễn Chí Thanh. Dựa vào nội dung file kế hoạch: {content_kh_thang}, soạn thảo nội dung phần II. NỘI DUNG CUỘC HỌP bóc tách lắp vào mục 1.1, 1.2, 1.3 phát triển chi tiết."
                        res_ai, _ = run_ai_prompt_safe(prompt_bb, api_key_input)
                        full_bien_ban_text = f"BIÊN BẢN SINH HOẠT TỔ CHUYÊN MÔN\\nKỲ HỌP THỨ {ky_hop_so} THÁNG {thang_hop_so}\\n\\nI. THỜI GIAN, ĐỊA ĐIỂM\\nLúc {gio_phut_hop}, ngày {ngay_thang_hop} tại {phong_hop}\\nChủ trì: {chu_tri_cuoc}, Thư ký: {thu_ky_cuoc}, Thành phần: {thanh_phan_co}\\n\\nII. NỘI DUNG CUỘC HỌP\\n{res_ai}\\n\\n2. Ý KIẾN ĐÓNG GÓP\\n{y_kien_dong_gop}"
                        st.markdown(full_bien_ban_text)
                        st.download_button("📥 Xuất văn bản biên bản Word (.docx)", data=export_to_docx_vietnam_standard(full_bien_ban_text, f"Bien_ban_thang_{thang_hop_so}"), file_name=f"Bien_ban_thang_{thang_hop_so}.docx")
                    except Exception as e: st.error(f"Lỗi: {e}")

    elif chuc_nang_chinh == "3. Xây dựng Kế hoạch Giáo dục cá nhân (Phụ lục III - Công văn 5512)":
        st.header("📋 Kế hoạch Giáo dục cá nhân của Giáo viên (Phụ lục III)")
        gv_selected = st.selectbox("Chọn Giáo viên cần lập kế hoạch:", [x["Họ và tên"] for x in st.session_state["db_thanh_vien"]])
        mon_tu_dong = next((x["Phân môn chính"] for x in st.session_state["db_thanh_vien"] if x["Họ và tên"] == gv_selected), "")
        st.info(f"Phân môn giảng dạy: **{mon_tu_dong}**")
        if st.button("🪄 AI Thiết kế Phụ lục III chuẩn 5512"):
            with st.spinner("AI đang xử lý..."):
                try:
                    res, _ = run_ai_prompt_safe(f"Thiết kế khung kế hoạch giáo dục cá nhân phụ lục III 5512 cho môn {mon_tu_dong}", api_key_input)
                    st.markdown(res)
                    st.download_button("📥 Tải Phụ lục III Word (.docx)", data=export_to_docx_vietnam_standard(res, f"Phu_luc_III_{gv_selected}"), file_name=f"Phu_luc_III_{gv_selected}.docx")
                except Exception as e: st.error(f"Lỗi: {e}")

    elif chuc_nang_chinh == "4. Thống kê số liệu tổ":
        st.header("📊 Trung tâm Thống kê Số liệu Tổ chuyên môn")
        tong_nhan_su = len(st.session_state["db_thanh_vien"])
        st.metric("Tổng số thành viên tổ", f"{tong_nhan_su} Giáo viên")
        df_tv_current = pd.DataFrame(st.session_state["db_thanh_vien"])
        if "Phân môn chính" in df_tv_current.columns:
            st.bar_chart(df_tv_current["Phân môn chính"].value_counts())
