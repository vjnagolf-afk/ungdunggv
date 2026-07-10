import streamlit as st
import gspread

# Cấu hình thông tin Google Sheets cố định theo ID và Tab của bạn
SPREADSHEET_ID = "1C6642jk_oQ0g9UC2By2qsNxxfQVR0MrZYj52tRdWDlY"
TAB_NAME = "DE_KT"

def sync_to_google_sheet(ten_de, mon, khoi, thoi_gian, noi_dung):
    """
    Hàm kết nối Google Sheets tích hợp thuật toán tự động dò tìm cấu hình chứng thực trong Secrets
    """
    try:
        creds_dict = None
        
        # 1. KIỂM TRA CÁC TÊN KHÓA ƯU TIÊN TRƯỚC
        priority_keys = ["gspread_credentials", "GSPREAD_CREDENTIALS", "google_sheet_creds", "google_creds", "GOOGLE_KEY"]
        for key in priority_keys:
            if key in st.secrets:
                creds_dict = st.secrets[key]
                break
                
        # 2. THUẬT TOÁN DÒ TÌM TỰ ĐỘNG NẾU KHÔNG TÌM THẤY TRONG CÁC KHÓA ƯU TIÊN
        if creds_dict is None:
            for key in st.secrets.keys():
                node = st.secrets[key]
                if hasattr(node, "get") or isinstance(node, dict):
                    if node.get("type") == "service_account" or "private_key" in node:
                        creds_dict = node
                        break
                    for sub_key in node.keys():
                        sub_node = node[sub_key]
                        if hasattr(sub_node, "get") or isinstance(sub_node, dict):
                            if sub_node.get("type") == "service_account" or "private_key" in sub_node:
                                creds_dict = sub_node
                                break
                if creds_dict: break

        if creds_dict is None:
            st.error("Hệ thống tự động quét toàn bộ Secrets nhưng không tìm thấy cấu hình Service Account hợp lệ nào!")
            return False
            
        gc = gspread.service_account_from_dict(creds_dict)
        sh = gc.open_by_key(SPREADSHEET_ID)
        worksheet = sh.worksheet(TAB_NAME)
        
        if len(worksheet.get_all_values()) == 0:
            worksheet.append_row(["Tên Đề", "Môn Học", "Khối Lớp", "Thời Gian", "Nội Dung Chi Tiết"])
            
        worksheet.append_row([ten_de, mon, khoi, thoi_gian, noi_dung])
        return True
        
    except Exception as e:
        st.warning(f"Không thể đồng bộ Google Sheet: {e}")
        return False
import io
import re
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from pypdf import PdfReader

def read_uploaded_docx(uploaded_file):
    """Đọc dữ liệu chữ từ file Word tải lên"""
    try:
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    except: 
        return "Lỗi đọc file Word"

def read_uploaded_pdf(uploaded_file):
    """Đọc dữ liệu chữ từ file PDF tải lên"""
    try:
        reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except: 
        return "Lỗi đọc file PDF"

def generate_plot_stream(eq_str):
    """Tự động biên dịch hàm số toán học và xuất luồng ảnh đồ thị"""
    fig, ax = plt.subplots(figsize=(5, 3.5))
    x = np.linspace(-10, 10, 400)
    
    safe_dict = {
        "x": x, "np": np, "sin": np.sin, "cos": np.cos, "tan": np.tan, "sqrt": np.sqrt
    }
    try:
        eq_str_py = eq_str.replace('^', '**')
        y = eval(eq_str_py, {"__builtins__": {}}, safe_dict)
        
        if isinstance(y, (int, float)):
            y = np.full_like(x, y)
        ax.plot(x, y, color='#1E40AF', linewidth=2.5)
        ax.axhline(0, color='black', linewidth=1.2)
        ax.axvline(0, color='black', linewidth=1.2)
        ax.grid(color='gray', linestyle='--', linewidth=0.5, alpha=0.7)
        
        ax.set_ylim([-10, 10])
        ax.set_title(f"Đồ thị: y = {eq_str}", fontsize=10, pad=10)
    except Exception as e:
        ax.text(0.5, 0.5, f"[Không thể vẽ đồ thị: Sai cú pháp toán học]", ha='center', va='center', color='red')
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def convert_latex_to_omml(latex_str):
    """Chuyển đổi các định dạng công thức Toán học LaTeX cơ bản sang XML Math của Microsoft Word"""
    latex_str = latex_str.replace(r'\pi', 'π').replace(r'\infty', '∞').replace(r'\times', '×')
    
    # Xử lý tổng xích-ma \sum_{k=0}^{n}
    latex_str = re.sub(r'\\sum_\{([^}]+)\}\^\{([^}]+)\}', r'∑(chạy từ \1 đến \2)', latex_str)
    
    # Xử lý phân số \frac{a}{b} -> (a)/(b)
    frac_pattern = re.compile(r'\\frac\{([^}]+)\}\{([^}]+)\}')
    while frac_pattern.search(latex_str):
        latex_str = frac_pattern.sub(r'(\1)/(\2)', latex_str)
        
    # Chuẩn hóa lũy thừa số mũ
    latex_str = re.sub(r'\^\{([^}]+)\}', r'^\1', latex_str)
    
    omml_xml = f'<w:p {nsdecls("w")}><m:oMathPara {nsdecls("m")}><m:oMath><m:r><m:t>{latex_str}</m:t></m:r></m:oMath></m:oMathPara></w:p>'
    try:
        return parse_xml(omml_xml)
    except:
        return None

def process_runs_with_math(paragraph, text):
    """Phân tách chuỗi chứa công thức toán trong cặp dấu $...$ và chữ thường để ghi vào Word"""
    parts = re.split(r'(\$\$.*?\$\$|\$.*?\$)', text)
    for part in parts:
        if part.startswith('$'): 
            math_content = part.replace('$$', '').replace('$', '')
            math_element = convert_latex_to_omml(math_content)
            if math_element is not None:
                paragraph._p.append(math_element)
            else:
                run = paragraph.add_run(part)
                run.font.name = 'Times New Roman'
        else:
            bold_parts = part.split('**')
            for i, b_part in enumerate(bold_parts):
                is_bold = (i % 2 != 0)
                sub_sup_parts = re.split(r'(<sub>.*?</sub>|<sup>.*?</sup>)', b_part)
                for s_part in sub_sup_parts:
                    if not s_part: continue
                    if s_part.startswith('<sub>') and s_part.endswith('</sub>'):
                        run = paragraph.add_run(s_part[5:-6])
                        run.bold = is_bold
                        run.font.subscript = True
                    elif s_part.startswith('<sup>') and s_part.endswith('</sup>'):
                        run = paragraph.add_run(s_part[5:-6])
                        run.bold = is_bold
                        run.font.superscript = True
                    else:
                        run = paragraph.add_run(s_part)
                        run.bold = is_bold
                        run.font.name = 'Times New Roman'

def export_to_docx_vietnam_standard(text_content, title_name, 
                                   school_name="TRƯỜNG THCS NGUYỄN CHÍ THANH", group_name="TỔ KHOA HỌC TỰ NHIÊN - GDTC"):
    """Dựng khung văn bản mẫu chuẩn hành chính Giáo dục Việt Nam"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    
    # Tạo tiêu đề bảng giám hiệu / tổ chuyên môn
    admin_table = doc.add_table(rows=1, cols=2)
    admin_table.autofit = False
    
    cell_l = admin_table.rows[0].cells[0]
    p_left = cell_l.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.add_run(f"{school_name.upper()}\n").bold = True
    p_left.add_run(f"{group_name.upper()}\n").bold = True
    p_left.add_run("Số: ..... /BB-TCM").font.size = Pt(11)
    
    cell_r = admin_table.rows[0].cells[1]
    p_right = cell_r.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").bold = True
    p_right.add_run("Độc lập - Tự do - Hạnh phúc\n").bold = True
    p_right.add_run("***************").font.size = Pt(11)
    
    admin_table.columns[0].width = Inches(3.2)
    admin_table.columns[1].width = Inches(3.8)
    
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run(title_name.upper()).bold = True
    p_title.runs[0].font.size = Pt(14)
    
    in_table = False
    table_data = []
    
    def build_table():
        if not table_data: return
        cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_val in enumerate(row):
                if c_idx < cols:
                    cell = table.cell(r_idx, c_idx)
                    p = cell.paragraphs[0]
                    p.text = "" 
                    process_runs_with_math(p, cell_val.strip())
        doc.add_paragraph() 
        
    for line in text_content.split('\n'):
        cleaned_line = line.strip()
        
        if cleaned_line.startswith('|') and cleaned_line.endswith('|'):
            in_table = True
            row_data = [cell.strip() for cell in cleaned_line.split('|')[1:-1]]
            if all(re.match(r'^[-: ]+$', cell) for cell in row_data): continue
            table_data.append(row_data)
            continue
        
        if in_table:
            build_table()
            in_table = False
            table_data = []
        
        if not cleaned_line: continue
        
        if '[GRAPH:' in cleaned_line:
            match = re.search(r'\[GRAPH:\s*(.+?)\]', cleaned_line)
            if match:
                eq = match.group(1)
                img_stream = generate_plot_stream(eq) # Gọi hàm từ Đoạn 2
                doc.add_picture(img_stream, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # 🚀 TỰ ĐỘNG PHÁT HIỆN VÀ ÉP XUỐNG DÒNG ĐÁP ÁN TRẮC NGHIỆM A, B, C, D
        if re.search(r'A\..*B\..*C\..*D\.', cleaned_line):
            sub_choices = re.split(r'(?=A\.|B\.|C\.|D\.)', cleaned_line)
            for choice in sub_choices:
                if choice.strip():
                    p = doc.add_paragraph()
                    process_runs_with_math(p, choice.strip())
            continue
            
        p = doc.add_paragraph()
        if re.match(r'^(Câu \d+:)', cleaned_line):
            prefix = re.match(r'^(Câu \d+:)', cleaned_line).group(1)
            rest = cleaned_line[len(prefix):]
            run_p = p.add_run(prefix + " ")
            run_p.bold = True
            process_runs_with_math(p, rest.strip())
        else:
            process_runs_with_math(p, cleaned_line)
            
    if in_table and table_data:
        build_table()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
import streamlit as st
from google import genai
from google.genai import types

def call_gemini_with_fallback(prompt_text, preferred_model):
    """
    Hàm gọi API Gemini có cơ chế tự động chuyển tiếp mô hình khi gặp lỗi hạn mức (Quota)
    """
    try:
        client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception as e:
        return f"❌ Lỗi cấu hình API Key trong Streamlit Secrets: {str(e)}", None

    # Thiết lập nhóm hạ cấp tương thích tốt nhất
    model_pool = {
        "3.1 Pro": ["gemini-1.5-pro", "gemini-1.5-flash", "gemini-2.5-flash-8b"],
        "3.5 Flash": ["gemini-1.5-flash", "gemini-2.5-flash-8b"],
        "Flash Lite": ["gemini-2.5-flash-8b"]
    }
    models_to_try = model_pool.get(preferred_model, ["gemini-1.5-flash"])
    
    last_error = ""
    for model_name in models_to_try:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=prompt_text,
                config=types.GenerateContentConfig(temperature=0.2)
            )
            return response.text, model_name
        except Exception as e:
            last_error = str(e)
            continue
            
    return f"❌ Tất cả mô hình đều lỗi: {last_error}", None
import streamlit as st

def main():
    st.set_page_config(page_title="Hệ thống sinh đề kiểm tra tự động", layout="wide")
    st.title("🎯 TRỢ LÝ THIẾT KẾ ĐỀ KIỂM TRA CHUẨN QUY ĐỊNH")
    
    if "ket_qua_de" not in st.session_state: st.session_state["ket_qua_de"] = ""
    if "model_dung" not in st.session_state: st.session_state["model_dung"] = ""

    # Giao diện tải File mẫu Ma trận / Đặc tả ở thanh bên trái
    st.sidebar.header("📁 CẤU HÌNH FILE MẪU & MA TRẬN")
    file_mau = st.sidebar.file_uploader("Tải lên file Ma trận hoặc Đặc tả mẫu (.docx, .pdf):", type=["docx", "pdf"])
    
    noi_dung_mau = ""
    if file_mau:
        if file_mau.name.endswith(".docx"):
            noi_dung_mau = read_uploaded_docx(file_mau) # Gọi hàm Đoạn 2
        else:
            noi_dung_mau = read_uploaded_pdf(file_mau)  # Gọi hàm Đoạn 2
        st.sidebar.success("Đã đọc nội dung cấu trúc file mẫu thành công!")

    if not noi_dung_mau:
        noi_dung_mau = st.sidebar.text_area("Hoặc dán yêu cầu Ma trận & Đặc tả cấu trúc đề vào đây:", 
                                            value="Chương 1: Hàm số (2 câu Nhận biết, 1 câu Thông hiểu)...\nĐáp án trắc nghiệm viết độc lập xuống dòng.")

    # Các thông tin tùy biến của đề kiểm tra
    col1, col2, col3 = st.columns(3)
    with col1:
        ten_de = st.text_input("Tên bài kiểm tra / Đề số:", value="Đề kiểm tra giữa học kỳ I")
        mon_hoc = st.text_input("Môn học:", value="Toán học")
    with col2:
        khoi_lop = st.text_input("Khối lớp:", value="Khối 12")
        thoi_gian = st.text_input("Thời gian làm bài:", value="90 phút")
    with col3:
        mo_hinh_ai = st.selectbox("Chọn mô hình AI ưu tiên:", ["3.1 Pro", "3.5 Flash", "Flash Lite"])
        school = st.text_input("Tên trường:", value="TRƯỜNG THCS NGUYỄN CHÍ THANH")

    if st.button("⚡ Tiến hành sinh đề thi tự động bằng AI", type="primary"):
        with st.spinner("🧠 AI đang phân tích ma trận mẫu và tạo đề toán học định dạng LaTeX..."):
            prompt_system = f"""
            Bạn là một chuyên gia ra đề thi học thuật cấp cao tại Việt Nam. 
            Nhiệm vụ của bạn là tạo một Đề kiểm tra môn {mon_hoc} cho {khoi_lop} với chủ đề: {ten_de}.
            
            QUY ĐỊNH BẮT BUỘC VỀ MA TRẬN VÀ CẤU TRÚC:
            Hãy bám sát cấu trúc phân phối, tỷ lệ nhận biết, thông hiểu, vận dụng từ dữ liệu đặc tả mẫu sau đây:
            {noi_dung_mau}
            
            QUY ĐỊNH NGHIÊM NGẶT VỀ ĐỊNH DẠNG:
            1. CÔNG THỨC TOÁN HỌC: Tất cả các biểu thức toán, phân số, số mũ, căn thức, tổng xích-ma, kí hiệu hình học, biến số phải được đặt trong cặp dấu đô-la đơn $...$ hoặc đôi $$. Ví dụ: $A = \\pi r^2$, $(x+a)^n = \\sum_{{k=0}}^{{n}} \\binom{{n}}{{k}} x^k a^{{n-k}}$, $f(x) = \\frac{{n}}{{1!}}$. Tuyệt đối không dùng chữ viết thường tự do.
            2. XUỐNG DÒNG PHƯƠNG ÁN: Các đáp án trắc nghiệm A., B., C., D. BẮT BUỘC mỗi đáp án phải nằm riêng biệt trên một dòng mới. Tuyệt đối không viết gộp trên cùng một hàng.
               Ví dụ:
               Câu 1: Khẳng định nào đúng?
               A. Đồ thị đồng biến
               B. Đồ thị nghịch biến
               C. Hàm số vô nghiệm
               D. Hàm số có 2 nghiệm
            """
            ket_qua, model_thuc_te = call_gemini_with_fallback(prompt_system, mo_hinh_ai) # Gọi hàm Đoạn 4
            st.session_state["ket_qua_de"] = ket_qua
            st.session_state["model_dung"] = model_thuc_te
            
            # Tự động đồng bộ thông tin lưu trữ lên Google Sheets khi sinh đề thành công
            if "❌" not in ket_qua:
                sync_to_google_sheet(ten_de, mon_hoc, khoi_lop, thoi_gian, ket_qua[:500] + "...") # Gọi hàm Đoạn 1

    # Khối kết xuất và tải tệp
    if st.session_state["ket_qua_de"]:
        if "❌" in st.session_state["ket_qua_de"]:
            st.error(st.session_state["ket_qua_de"])
        else:
            st.info(f"🤖 Đề thi được tối ưu và khởi tạo thành công bởi mô hình: `{st.session_state['model_dung']}`")
            st.markdown("### 📝 Xem trước đề thi:")
            st.markdown(st.session_state["ket_qua_de"])
            
            # Xuất file Word chuẩn
            data_word = export_to_docx_vietnam_standard(st.session_state["ket_qua_de"], ten_de, school_name=school) # Gọi hàm Đoạn 3
            
            st.download_button(
                label="📥 Tải tệp Đề kiểm tra Word (.docx) bản chuẩn hành chính",
                data=data_word,
                file_name=f"De_Kiem_Tra_{khoi_lop.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

if __name__ == "__main__":
    main()
